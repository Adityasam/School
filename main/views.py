from django.shortcuts import render, HttpResponse, redirect
from .models import *
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from .helper import *
from django.core.files.storage import FileSystemStorage
from django.conf import settings as django_settings
import os
import base64
from django.db import connection
import docx
import datetime
from django.utils import timezone
from docx2pdf import convert
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime, time
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import calendar as newcalendar
import PIL.Image as Image
import io
from django.db.models.signals import post_save
from main.face_rec import *
import numpy as np
from .messenger import *

# Create your views here.

project_folder_path = django_settings.BASE_DIR

def detect_device(request):
    is_mobile = False
    user_agent = request.META.get('HTTP_USER_AGENT', '').lower()
    mobile_keywords = ['iphone', 'android', 'windows phone']
    for keyword in mobile_keywords:
        if keyword in user_agent:
            is_mobile = True
            break

    return is_mobile

def face_test(request):
    return render(request, "face_test.html")

@csrf_exempt
def clear_old_face(request):
    code = request.POST.get("usercode")
    print(code)
    delete_old_images(code)

    return JsonResponse({"status":"success"})

@csrf_exempt
def encrypt_face(request):
    image = request.FILES["file"].read()
    empcode = request.POST.get("code")
    image = Image.open(io.BytesIO(image))
    rgb_im = image.convert('RGB')
    im_src_array=np.array(rgb_im)
    result = encode_face(im_src_array, empcode)

    create_user(empcode)

    return JsonResponse({"status":result})

def create_user(code):
    code = code.split("^")[0]
    type = ""
    if code[:3] == "STD":
        type = "P"
    elif code[:3] == "FAC":
        type = "F"
    elif code[:3] == "STF":
        type = "S"

    old = MyUser.objects.filter(UserId = code)
    if len(old) == 0:
        newuser = MyUser()
        newuser.UserId = code
        newuser.Password = "8896255889"
        newuser.UserType = type
        newuser.save()

        #Giving Default Permissions
        defs = RolePrivilege.objects.filter(RoleCode = type)
        for d in defs:
            u = UserPrivilege()
            u.UserId = code
            u.PrivilegeCode = d.PrivilegeCode
            u.save()

@csrf_exempt
def recognize_face(request):
    result = {}
    
    image = request.FILES["file"].read()
    image = Image.open(io.BytesIO(image))
    rgb_im = image.convert('RGB')
    im_src_array=np.array(rgb_im)
    result = recognize_face_image(im_src_array)
    print(result["result"])

    return JsonResponse(result, status=201)

def permission(request):
    if check_login(request) == False or not request.session["SUPERADMIN"]:
        return redirect("main:login")
    
    fac = []
    stf = []
    
    role = Role.objects.filter(Active = True, Deleted = False)
    role = parse_data(role)
    
    with connection.cursor() as cursor:
        cursor.execute("""select f.FacultyCode as usercode, f.FullName as FullName, 'F' as Type from main_Faculty f""")
        fac = parse_curser(cursor)

    with connection.cursor() as cursor:
        cursor.execute("""select f.StaffCode as usercode, f.FullName as FullName, 'S' as Type from main_Staff f""")
        stf = parse_curser(cursor)

    perms = Menu.objects.filter(Visible = True).order_by("-SortOrder")
    perms = parse_data(perms) 
            
    context = {
        "employees": fac+stf,
        "perms":perms,
        "roles":role
    }
    return render(request, "permission.html", context)

def roles(request):
    role = []
    with connection.cursor() as cursor:
        cursor.execute("""select r.* from main_Role r""")
        role = parse_curser(cursor)

    perms = Menu.objects.filter(Visible = True).order_by("-SortOrder")
    perms = parse_data(perms) 
            
    context = {
        "role": role,
        "perms":perms
    }
    return render(request, "roles.html", context)

def check_login(request):
    if "USERID" not in request.session:
        return False
    else:
        return True

def receipt_template(request, id):
    alldata = []
    items = []
    with connection.cursor() as cursor:
        cursor.execute("""select f.*, s.FullName
        from main_Fee f
        left join main_Student s on s.StudentCode = f.StudentCode
        where f.id=%s""", [id])

        alldata = parse_curser(cursor)

    with connection.cursor() as cursor:
        cursor.execute("""select f.*, ft.Title as feetypetitle 
        from main_FeeItem f
        left join main_FeeType ft on ft.id = f.Type
        where f.FeeId=%s""", [id])

        items = parse_curser(cursor)

    return render(request, "template/receipt_template.html", {"data":alldata, "items":items})

def mark_sheet(request):
    return render(request, "template/mark_sheet.html")

def attendance(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00011"
    classes = Class.objects.filter(Deleted = False).order_by("Title")
    sections = Section.objects.filter(Deleted = False)
    sessions = SchoolSession.objects.filter(Deleted = False)

    context = {
        'classes': classes,
        'sections' : sections,
        'sessions' : sessions
    }
    return render(request, "attendance.html", context)

def facultyattendance(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00014"
    sessions = SchoolSession.objects.filter(Deleted = False)

    context = {
        'sessions' : sessions
    }
    return render(request, "facultyattendance.html", context)

def salary(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00013"
    sessions = SchoolSession.objects.filter(Deleted = False)
    context = {
        'sessions': sessions
    }
    return render(request, "salary.html", context)

def announcement(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00037"
    return render(request, "announcement.html")

def mobile_notice(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00038"
    return render(request, "mobile/mobile_notice.html", {"pagename": "Announcement"})
    
def syllabus(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00012"
    classes = Class.objects.filter(Deleted = False).order_by("Title")
    sections = Section.objects.filter(Deleted = False)
    sessions = SchoolSession.objects.filter(Deleted = False)
    subject = Subject.objects.filter(Deleted = False)

    context = {
        'classes': classes,
        'sections' : sections,
        'sessions' : sessions,
        'subjects' : subject
    }
    return render(request, "syllabus.html", context)

def dashboard(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00009"
    return render(request, "dashboard.html")

def library_dashboard(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00046"
    return render(request, "library_dashboard.html")

def member(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    batches = Batch.objects.filter(Deleted = False, CompanyCode = request.session["COMPANY"])
    batches = parse_data(batches)
    context = {
        'batches': batches
    }
    request.session["PAGECODE"] = "PG_00039"
    return render(request, "member.html", context)

def member_attendance(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    batches = Batch.objects.filter(Deleted = False, CompanyCode = request.session["COMPANY"])
    batches = parse_data(batches)
    context = {
        'batches': batches
    }
    request.session["PAGECODE"] = "PG_00040"
    return render(request, "member_attendance.html", context)

def staff_list(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00021"
    
    design = Designation.objects.filter(Deleted = False, Active = True)
    context = {
        "designations":design
    }
    
    return render(request, "staff_list.html", context)

def login(request):
    if "HOMEPAGE" in request.session:
        return redirect("main:"+request.session["HOMEPAGE"])
    if detect_device(request) == True:
        return render(request, "mobile/mobilelogin.html")
    else:
        return render(request, "login.html")

def logmeout(request):
    request.session.clear()
    return redirect("main:login")

@csrf_exempt
def load_roles(request):
    role = Role.objects.filter(Active = True, Deleted = False)
    role = parse_data(role)
    return JsonResponse(role, safe=False)

@csrf_exempt
def get_user_notice(request):
    userid = request.POST.get("userid")
    notice = []
    with connection.cursor() as cursor:
        cursor.execute("""select n.* 
        from main_NoticeUser u
        inner join main_Notice n on n.id = u.noticeid
        where u.UserCode=%s order by senddate desc""", [userid])
        notice = parse_curser(cursor)

    return JsonResponse(notice, safe=False)

@csrf_exempt
def get_message_token(request):
    userid = request.POST.get("usercode")
    token = MessageId.objects.filter(UserId = userid)
    token = parse_data(token)

    return JsonResponse(token, safe=False) 

@csrf_exempt
def get_announcements(request):
    announce = []
    with connection.cursor() as cursor:
        cursor.execute("""select n.*,
        (case when n.status = 'N' then 'Saved' when n.status = 'A' then 'Sent' end) as statustitle 
        from main_Notice n""")
        announce = parse_curser(cursor)

    return JsonResponse(announce, safe=False)

@csrf_exempt
def get_notice_detail(request):
    noticeid = request.POST.get("noticeid")
    notice = Notice.objects.filter(id = noticeid)
    notice = parse_data(notice)[0]

    return JsonResponse(notice, safe=False)

@csrf_exempt
def load_users(request):
    users = []
    with connection.cursor() as cursor:
        cursor.execute("""select u.*, r.title as roletitle from main_MyUser u
        left join main_Role r on r.rolecode = u.usertype
        where u.Deleted = 0""")

        users = parse_curser(cursor)
    return JsonResponse(users, safe=False)

@csrf_exempt
def face_login(request):
    userid = request.POST.get("usercode")
    request.session["USERID"] = userid
    if userid[:3] == "STD":
        request.session["USERTYPE"] = "P"
    request.session["LOGINFROM"] = "M"

@csrf_exempt
def get_face_data(request):
    type = request.POST.get("type")
    
    faces = []
    if type == "F":
        fac = Faculty.objects.filter(Deleted = 0, Active = 1)
        fac = parse_data(fac)
        for f in fac:
            temp = {}
            code = f["facultycode"]
            pics = get_user_photos(code)
            temp["usercode"] = code
            temp["photo"] = f["photo"]
            temp["fullname"] = f["fullname"]
            temp["images"] = pics

            faces.append(temp)
    
    if type == "P":
        fac = Student.objects.filter(Deleted = 0, Active = 1)
        fac = parse_data(fac)
        for f in fac:
            temp = {}
            code = f["studentcode"]
            pics = get_user_photos(code)
            temp["usercode"] = code
            temp["photo"] = f["photo"]
            temp["fullname"] = f["fullname"]
            temp["images"] = pics

            faces.append(temp)

    if type == "S":
        fac = Staff.objects.filter(Deleted = 0, Active = 1)
        fac = parse_data(fac)
        for f in fac:
            temp = {}
            code = f["staffcode"]
            pics = get_user_photos(code)
            temp["usercode"] = code
            temp["photo"] = f["photo"]
            temp["fullname"] = f["fullname"]
            temp["images"] = pics

            faces.append(temp)

    return JsonResponse(faces, safe=False)

@csrf_exempt
def VerifyLogin(request):
    loginid = request.POST.get("loginid")
    password = request.POST.get("password")
    loginfrom = request.POST.get("from")
    logintype = request.POST.get("logintype") #F for facelogin, N for normal login
    if logintype == "":
        logintype = "N"
    usr = []
    if logintype == 'N':
        usr = MyUser.objects.filter(UserId__iexact = loginid, Password = password)
    else:
        usr = MyUser.objects.filter(UserId__iexact = loginid)

    if len(usr) > 0:
        usr = usr[0]
        if usr.Blocked:
            return JsonResponse({"status":"BLOCKED"})
        else:
            request.session["USERID"] = usr.UserId
            request.session["SUPERADMIN"] = usr.IsSuperAdmin
            request.session["EMPCODE"] = usr.EmployeeCode
            request.session["LOGINFROM"] = loginfrom
            request.session["USERTYPE"] = usr.UserType
            request.session["COMPANY"] = usr.CompanyCode

            menu = []
            homepage = ""
            if request.session["SUPERADMIN"] == True:
                with connection.cursor() as cursor:
                    cursor.execute("""select m.* from main_Menu m
                    where ((m.Desktop=1 and 'D'=%s) or (m.Mobile=1 and 'M'=%s)) and m.Visible = 1 and m.ForAdmin = 1
                    order by SortOrder desc""", 
                    [request.session["LOGINFROM"], request.session["LOGINFROM"]])
                    menu = parse_curser(cursor)
            else:
                with connection.cursor() as cursor:
                    cursor.execute("""select m.* from main_Menu m
                    inner join main_UserPrivilege p on p.UserId = %s and m.PageCode = p.PrivilegeCode 
                    where ((m.Desktop=1 and 'D'=%s) or (m.Mobile=1 and 'M'=%s))  and m.Visible = 1
                    order by SortOrder desc""", 
                    [request.session["USERID"], request.session["LOGINFROM"], request.session["LOGINFROM"]])
                    menu = parse_curser(cursor)

            if len(menu) > 0:
                homepage = menu[0]["pageurl"]
            
            request.session["HOMEPAGE"] = homepage

            perm = UserPrivilege.objects.filter(UserId = usr.UserId)
            perm = parse_data(perm)
            allcodes = []
            for p in perm:
                allcodes.append(p["privilegecode"])

            allcodes = "^".join(allcodes)
            
            request.session["PRIVILEGE"] = allcodes
            return JsonResponse({"status":"SUCCESS", "homepage":homepage})
    else:
        return JsonResponse({"status":"FAILED"})

def faculty_list(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00001"
    subject = Subject.objects.filter(Deleted = False)
    context = {
        'subjects': subject
    }
    return render(request, "faculty.html", context)

def report_card(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00007"
    classes = Class.objects.filter(Deleted = False).order_by("Title")
    sections = Section.objects.filter(Deleted = False)
    exams = None
    with connection.cursor() as cursor:
        cursor.execute("""select e.*, s.Title as SessionTitle from main_Exam e
        left join main_SchoolSession s on s.id = e.SessionId
        where e.Deleted = 0""")
        exams = parse_curser(cursor)

    context = {
        'classes': classes,
        'sections' : sections,
        'exams' : exams
    }
    return render(request, "report_card.html", context)

def paper_list(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00006"
    exams = None
    papers = None
    with connection.cursor() as cursor:
        cursor.execute("""select e.*, s.Title as SessionTitle from main_Exam e
        left join main_SchoolSession s on s.id = e.Sessionid
        where e.Deleted = 0 and e.enddate > CURRENT_DATE""")
        exams = parse_curser(cursor)

    with connection.cursor() as cursor:
        cursor.execute("""select e.*, x.Title as examtitle, s.Title as subject from main_ExamPaper e
        left join main_Exam x on x.id = e.ExamId
        left join main_Subject s on s.id = e.SubjectId
        where e.Deleted = 0""")
        papers = parse_curser(cursor)

    cls = Class.objects.filter(Deleted = False).order_by("Title")
    sections = Section.objects.filter(Deleted=False)
    subjects = Subject.objects.filter(Deleted = False)
    context = {
        'exams': exams,
        'classes' : cls,
        'subjects': subjects,
        'sections' : sections,
        'papers' : papers
    }

    return render(request, "paper_list.html", context)


def exam_list(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00004"
    sessions = SchoolSession.objects.filter(Deleted = False)
    context = {
        'sessions':sessions
    }
    return render(request, "exam_list.html", context)

def class_routine(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00003"
    subject = Subject.objects.filter(Deleted = False)
    faculty = Faculty.objects.filter(Deleted = False)
    cls = Class.objects.filter(Deleted = False).order_by("Title")
    section = Section.objects.filter(Deleted = False)
    sessions = SchoolSession.objects.filter(Deleted = False)
    context = {
        'classes': cls,
        'sections' : section,
        'subject' : subject,
        'faculty' : faculty,
        'sessions' : sessions
    }

    return render(request, "class_routine.html", context)

def faculty_routine(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00015"
    faculty = Faculty.objects.filter(Deleted = False)
    sessions = SchoolSession.objects.filter(Deleted = False)
    context = {
        'faculty' : faculty,
        'sessions' : sessions
    }

    return render(request, "faculty_routine.html", context)

def faculty_dashboard_mobile(request):
    if check_login(request) == False:
        return redirect("main:login")

    request.session["PAGECODE"] = "PG_00048"
    return render(request, "mobile/faculty_dashboard_mobile.html", {"pagename":"Dashboard"})

def member_dashboard_mobile(request):
    if check_login(request) == False:
        return redirect("main:login")

    request.session["PAGECODE"] = "PG_00016"
    return render(request, "mobile/member_dashboard_mobile.html", {"pagename":"Dashboard"})

def faculty_profile(request, faculty_id=""):
    if check_login(request) == False:
        return redirect("main:login")

    request.session["PAGECODE"] = "PG_00034"

    context = {
        "pagename" : "Profile",
        "facultycode" : faculty_id
    }

    return render(request, "mobile/faculty_profile.html", context)

def faculty_routine_mobile(request):
    if check_login(request) == False:
        return redirect("main:login")

    request.session["PAGECODE"] = "PG_00036"

    context = {
        "pagename" : "Routine"
    }

    return render(request, "mobile/faculty_routine.html", context)

def student_profile(request, student_id=""):
    if check_login(request) == False:
        return redirect("main:login")

    request.session["PAGECODE"] = "PG_00035"

    context = {
        "pagename" : "Profile",
        "student_id":student_id
    }

    return render(request, "mobile/student_profile.html", context)

def member_profile(request, member_id=""):
    if check_login(request) == False:
        return redirect("main:login")

    request.session["PAGECODE"] = "PG_00043"

    context = {
        "pagename" : "Profile",
        "member_id": member_id
    }

    return render(request, "mobile/member_profile.html", context)

def discussion_board(request):
    if check_login(request) == False:
        return redirect("main:login")

    request.session["PAGECODE"] = "PG_00029"
    return render(request, "mobile/discussion_board.html", {"pagename":"Discussion"})

def member_fee_collection(request):
    if check_login(request) == False:
        return redirect("main:login")

    members = LibraryMember.objects.filter(Deleted = False, CompanyCode = request.session["COMPANY"])
    members = parse_data(members)

    mode = PaymentMode.objects.filter(Deleted = False)
    mode = parse_data(mode)
    feetype = FeeType.objects.filter(Deleted = False)
    feetype = parse_data(feetype)
    batches = Batch.objects.filter(CompanyCode = request.session["COMPANY"])

    context = {
        'members' : members,
        'modes':mode,
        'feetype':feetype,
        'batches': batches
    }

    request.session["PAGECODE"] = "PG_00041"
    return render(request, "member_fee_collection.html", context)


def attendancemobile(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00017"
    cls = Class.objects.filter(Deleted = False).order_by("Title")
    sections = Section.objects.filter(Deleted = False)
    
    context={
        'classes': cls,
        "pagename": "Attendance",
        "sections": sections
    }
    return render(request, "mobile/mobileattendance.html", context)

def mobileattendance_member(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00045"
    batch = Batch.objects.filter(Deleted = False, CompanyCode = request.session["COMPANY"])
    
    context={
        'batch': batch,
        "pagename": "Attendance"
    }
    return render(request, "mobile/mobileattendance_member.html", context)

def dashboard_mobile(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00018"
    return render(request, "mobile/dashboard_mobile.html", {"pagename":"Dashboard"})


def library_dashboard_mobile(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00047"
    return render(request, "mobile/library_dashboard_mobile.html", {"pagename":"Dashboard"})

def parent_dashboard_mobile(request):
    request.session["PAGECODE"] = "PG_00022"
    studentcode = request.session["USERID"]
    student = []
    attendance = []
    ses = SchoolSession.objects.filter(Deleted = False, Active=True)
    current_session = ses[0].id

    currentmonth = timezone.now().strftime("%Y-%m")
    currentday = timezone.now().strftime("%d")

    num_days = newcalendar.monthrange(timezone.now().year, timezone.now().month)[1]

    classid = ""
    with connection.cursor() as cursor:
        cursor.execute("""select s.*, c.Title as classtitle, sc.Title as sectiontitle from main_Student s
        left join main_Class c on c.id = s.Class
        left join main_Section sc on sc.id = s.Section
        where s.StudentCode=%s""", [studentcode])
        
        student = parse_curser(cursor)
        classid = student[0]["class"]

    with connection.cursor() as cursor:
        cursor.execute("""select coalesce(SUM(case when a.status='A' then 1 else 0 end), 0) as AbsentCount from main_Attendance a
        where a.StudentCode = %s and strftime('%%Y-%%m', a.AttendanceDate) = %s""", [studentcode, currentmonth])
        
        attendance = parse_curser(cursor)

    syllabus = []
    with connection.cursor() as cursor:
        cursor.execute("""select Distinct c.*,
        (select COUNT(s.id) from main_Syllabus s where c.id = s.ClassId and s.SessionId = %s and
        (select COUNT(sb.id) from main_Syllabus sb where sb.parentid = s.id) = 0) as totalchapter,
        (select COUNT(s.id) from main_Syllabus s where c.id = s.ClassId and s.SessionId = %s and s.Completed = 1 and
        (select COUNT(sb.id) from main_Syllabus sb where sb.parentid = s.id) = 0) as donechapter
        from main_Class c 
        inner join main_Student std on std.Class = c.id
        where std.StudentCode = %s""", [current_session, current_session, studentcode])
        syllabus = parse_curser(cursor)

    marks = []
    with connection.cursor() as cursor:
        cursor.execute("""select e.Title as ExamTitle,
        coalesce((select SUM(totalmarks) from main_ExamPaper p 
        inner join main_Class c on c.id = p.ClassId
        where p.ExamId = e.id and p.ClassId = %s),0) as totalmarks,
        coalesce((select SUM(marks) from main_ReportCard rc
        inner join main_ExamPaper ep on ep.id = rc.PaperId
        where rc.ExamId = e.id and rc.StudentCOde = %s),0) as marksobtained
        from main_Exam e
        where e.SessionId=%s""", [classid, studentcode,current_session])

        marks = parse_curser(cursor)
        allmarks = []
        for m in marks:
            tmp = {}
            if(float(m["totalmarks"])==0):
                tmp["percent"] = 0
            else:
                per = float(m["marksobtained"])/float(m["totalmarks"])
                tmp["percent"] = format(per*100, ".2f")
            tmp["title"] = m["examtitle"]
            allmarks.append(tmp)

    payments = []
    with connection.cursor() as cursor:
        cursor.execute("""select fi.*, f.CreatedDate, ft.Title as FeeTypeTitle 
        from main_Fee f 
        left join main_FeeItem fi on fi.FeeId = f.id
        left join main_FeeType ft on ft.id = fi.Type
        where f.StudentCode=%s""", [studentcode])

        payments = parse_curser(cursor)

    absent = attendance[0]["absentcount"]

    totalchapter = syllabus[0]["totalchapter"]
    donechapter = syllabus[0]["donechapter"]

    doneper = 0
    if totalchapter > 0:
        doneper = int((donechapter/totalchapter)*100)
    pendingper = 100-doneper

    if len(payments) > 3:
        payments = payments[:3]

    remarks = get_faculty_remarks(request.session["USERID"])

    context={
        "student":student,
        "pagename": "Dashboard",
        "absent":absent,
        "present":int(currentday)-absent,
        "donesyllabus":doneper,
        "pendingsyllabus":pendingper,
        "marks": allmarks,
        "payments":payments,
        "remarks": remarks
    }

    return render(request, "mobile/parent_dashboard_mobile.html", context)

def mobilesyllabus(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00019"
    cls = Class.objects.filter(Deleted = False).order_by("Title")
    sub = Subject.objects.filter(Deleted=False)
    sections = Section.objects.filter(Deleted = False)

    context={
        'classes': cls,
        'subjects' : sub,
        "pagename": "Syllabus",
        "sections": sections
    }
    return render(request, "mobile/mobilesyllabus.html", context)

def expense(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00010"
    type = ExpenseType.objects.filter(Deleted = False)
    mode = PaymentMode.objects.filter(Deleted = False)
    sessions = SchoolSession.objects.filter(Deleted = False)
    context = {
        'type': type,
        'modes' : mode,
        'sessions' : sessions
    }
    return render(request, "expense.html", context)

def student_list(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00002"
    classes = Class.objects.filter(Deleted = False).order_by("Title")
    sections = Section.objects.filter(Deleted = False)
    bus = Bus.objects.filter(Deleted = False)
    context = {
        'classes': classes,
        'sections' : sections,
        'busses':bus
    }
    return render(request, "student.html", context)

def settings(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00008"
    faculty = Faculty.objects.filter(Deleted = False)
    faculty = parse_data(faculty)
    cls = Class.objects.filter(Deleted=False).order_by("Title")
    subject = Subject.objects.filter(Deleted = False)
    sess = SchoolSession.objects.filter(Deleted = False)
    context = {
        'faculty': faculty,
        'classes': cls,
        'subjects': subject,
        'sessions' : sess
    }
    return render(request, "settings.html", context)

def fee_collection(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00005"
    student = Student.objects.filter(Deleted = False)
    mode = PaymentMode.objects.filter(Deleted = False)
    feetype = FeeType.objects.filter(Deleted = False)
    sessions = SchoolSession.objects.filter(Deleted = False)
    context = {
        'students': student,
        'modes': mode,
        'feetype': feetype,
        'sessions' : sessions
    }
    return render(request, "fee_collection.html", context)

def calendar(request):
    request.session["PAGECODE"] = "PG_00049"
    context = {
        'colrange':range(7),
        'rowrange':range(6)
    }
    return render(request, "calendar.html", context)

def get_faculty_remarks(receivercode=""):
    remarks = []
    with connection.cursor() as cursor:
        cursor.execute("""select r.*, f.FullName as FacultyName
        from main_FacultyRemark r
        left join main_MyUser mu on mu.UserId = r.Sender and mu.UserType = 'F'
        left join main_Faculty f on f.FacultyCode = mu.UserId
        where (r.receiver = %s) and r.Read = 0""", 
        [receivercode])

        remarks = parse_curser(cursor)

    return remarks

def parent_attendance(request):
    request.session["PAGECODE"] = "PG_00023"
    return render(request, "mobile/parent_attendance.html", {"pagename":"Attendance"})

def parent_syllabus(request):
    return render(request, "mobile/parent_syllabus.html", {"pagename":"Syllabus"})

def home_work(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00020"

    sessions = SchoolSession.objects.filter(Deleted=False)
    classes = Class.objects.filter(Deleted = False)
    subjects = Subject.objects.filter(Deleted = False)

    context = {
        "sessions":sessions,
        "classes":classes,
        "subjects":subjects
    }

    return render(request, "home_work.html", context)

def faculty_homework_mobile(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00027"

    sessions = SchoolSession.objects.filter(Deleted=False)
    classes = Class.objects.filter(Deleted = False)
    subjects = Subject.objects.filter(Deleted = False)

    context = {
        "sessions":sessions,
        "classes":classes,
        "subjects":subjects,
        "pagename":"Home Work"
    }

    return render(request, "mobile/faculty_homework_mobile.html", context)

def parent_homework(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00028"

    context = {
        "pagename": "Home Work"
    }

    return render(request, "mobile/parent_homework.html", context)
    
def mobile_faculty_list(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00030"
    context = {
        "pagename": "Faculty List"
    }
    return render(request, "mobile/mobile_faculty_list.html", context)

def mobile_student_list(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00031"
    classes = Class.objects.filter(Deleted = False)
    context = {
        "pagename" : "Student List",
        "classes" : classes
    }
    return render(request, "mobile/mobile_student_list.html", context)

def mobile_member_list(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00042"
    batches = Batch.objects.filter(Deleted = False, CompanyCode = request.session["COMPANY"])
    context = {
        "pagename" : "Member List",
        "batches" : batches
    }
    return render(request, "mobile/mobile_member_list.html", context)


def exam_paper_mobile(request):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00032"
    exams = None
    papers = None
    with connection.cursor() as cursor:
        cursor.execute("""select e.*, s.Title as SessionTitle from main_Exam e
        left join main_SchoolSession s on s.id = e.Sessionid
        where e.Deleted = 0""")
        exams = parse_curser(cursor)

    with connection.cursor() as cursor:
        cursor.execute("""select e.*, x.Title as examtitle, s.Title as subject from main_ExamPaper e
        left join main_Exam x on x.id = e.ExamId
        left join main_Subject s on s.id = e.SubjectId
        where e.Deleted = 0""")
        papers = parse_curser(cursor)

    cls = Class.objects.filter(Deleted = False).order_by("Title")
    sections = Section.objects.filter(Deleted=False)
    subjects = Subject.objects.filter(Deleted = False)
    context = {
        'exams': exams,
        'classes' : cls,
        'subjects': subjects,
        'sections' : sections,
        'papers' : papers,
        "pagename": "Exam Paper"
    }
    return render(request, "mobile/exam_paper_mobile.html", context)

def mobileattendance_faculty(request, type="F"):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00033"
    context = {
        "pagename" : "Faculty Attendance",
        "type" : type
    }
    return render(request, "mobile/mobileattendance_faculty.html", context)

def discussion_page(request, otherparty=""):
    if check_login(request) == False:
        return redirect("main:login")
    
    request.session["PAGECODE"] = "PG_00029"
    with connection.cursor() as cursor:
        cursor.execute("""update main_FacultyRemark set Read = 1 where Receiver = %s""", [request.session["USERID"]])
    
    user = MyUser.objects.filter(UserId = otherparty)
    
    rname = ""
    if len(user) > 0:
        user = user[0]
        type = user.UserType
        if type == "F":
            fcode = user.UserId
            faculty = Faculty.objects.filter(FacultyCode = fcode)
            faculty = faculty[0]
            rname = faculty.FullName
        if type == "P":
            fcode = user.UserId
            faculty = Student.objects.filter(StudentCode = fcode)
            faculty = faculty[0]
            rname = faculty.FullName
    
    if rname == "":
        rname = "Discussion"
    return render(request, "mobile/discussion_page.html", {"pagename":rname, "receiver": otherparty})
    

def parent_payment(request):
    request.session["PAGECODE"] = "PG_00024"
    sessions = SchoolSession.objects.filter(Deleted=False)

    context = {
        "sessions":sessions,
        "pagename":"Payments"
    }
    return render(request, "mobile/parent_payment.html", context)

def mobile_fee_collection(request):
    request.session["PAGECODE"] = "PG_00044"
    members = LibraryMember.objects.filter(Deleted = False, CompanyCode = request.session["COMPANY"])
    members = parse_data(members)

    mode = PaymentMode.objects.filter(Deleted = False)
    mode = parse_data(mode)
    feetype = FeeType.objects.filter(Deleted = False)
    feetype = parse_data(feetype)

    batches = Batch.objects.filter(CompanyCode = request.session["COMPANY"], Deleted= False)

    context = {
        "pagename":"Payments",
        "members": members,
        "modes":mode,
        "feetype":feetype,
        "batches":batches
    }
    return render(request, "mobile/mobile_fee_collection.html", context)

def parent_report_card(request):
    request.session["PAGECODE"] = "PG_00025"
    sessions = SchoolSession.objects.filter(Deleted=False)

    context = {
        "sessions":sessions,
        "pagename":"Report Card"
    }
    return render(request, "mobile/parent_report_card.html", context)

@csrf_exempt
def get_faculty_list_attendance(request):
    date = request.POST.get("date")

    faculty = []
    with connection.cursor() as cursor:
        cursor.execute("""select f.facultycode,f.photo, f.fullname, (case when a.status = 'P' then 'P' else 'A' end) as status from main_Faculty f
        left join main_FacultyAttendance a on a.FacultyCode = f.FAcultyCode 
        and strftime('%%Y-%%m-%%d', a.AttendanceDate) = %s""", [date])

        faculty = parse_curser(cursor)

    return JsonResponse(faculty, safe=False)

@csrf_exempt
def get_member_list_attendance(request):
    date = request.POST.get("date")
    batch = request.POST.get("batch")
    comp = request.session["COMPANY"]

    faculty = []
    with connection.cursor() as cursor:
        cursor.execute("""select f.membercode,f.picture, f.fullname, 
        (case when a.status = 'P' then 'P' else 'A' end) as status from main_LibraryMember f
        left join main_Attendance a on a.StudentCode = f.MemberCode 
        and strftime('%%Y-%%m-%%d', a.AttendanceDate) = %s
        where (%s='' or f.Batchid = %s) and (%s='' or f.CompanyCode=%s) order by f.FullName""", [date, batch, batch, comp, comp])

        faculty = parse_curser(cursor)

    return JsonResponse(faculty, safe=False)

@csrf_exempt
def get_staff_list_attendance(request):
    date = request.POST.get("date")

    staff = []
    with connection.cursor() as cursor:
        cursor.execute("""select f.staffcode,f.photo, f.fullname, (case when a.status = 'P' then 'P' else 'A' end) as status 
        from main_Staff f
        left join main_FacultyAttendance a on a.FacultyCode = f.StaffCode 
        and strftime('%%Y-%%m-%%d', a.AttendanceDate) = %s""", [date])

        staff = parse_curser(cursor)

    return JsonResponse(staff, safe=False)

@csrf_exempt
def send_message(request):
    message = request.POST.get("message")
    receiver = request.POST.get("receiver")
    sender = request.POST.get("sender")

    fr = FacultyRemark()
    fr.Sender = sender
    fr.Receiver = receiver
    fr.Message = message
    fr.save()

    return JsonResponse({"status":"success"})

@csrf_exempt
def send_notice(request):
    noticeid = request.POST.get("noticeid")
    title = request.POST.get("title")
    desc = request.POST.get("description")
    status = request.POST.get("status")
    sendto = request.POST.get("sendto")
    sentby = request.POST.get("sentby")

    notice = Notice.objects.filter(id = noticeid)
    if len(notice) > 0:
        notice = notice[0]
    else:
        notice = Notice()

    notice.Title = title
    notice.Description = desc
    notice.Status = status
    notice.SendTo = sendto
    notice.CreatedBy = sentby
    if status == 'A':
        notice.SendDate = timezone.now()
    notice.save()

    nid = notice.id


    sender = {}
    sender["title"] = title
    sender["message"] = desc

    tokens = []
    users = []
    if status == 'A':
        if "P" in sendto:
            students = Student.objects.filter(Deleted = False, Active = True)
            for s in students:
                sn = NoticeUser()
                sn.NoticeId = nid
                sn.UserCode = s.StudentCode
                sn.save()
                mid = MessageId.objects.filter(UserId = s.StudentCode)
                if len(mid) > 0:
                    mid = mid[0]
                    tokens.append(mid.Token)
                users.append(s.StudentCode)

        if "F" in sendto:
            faculty = Faculty.objects.filter(Deleted = False, Active = True)
            for s in faculty:
                sn = NoticeUser()
                sn.NoticeId = nid
                sn.UserCode = s.FacultyCode
                sn.save()
                mid = MessageId.objects.filter(UserId = s.FacultyCode)
                if len(mid) > 0:
                    mid = mid[0]
                    tokens.append(mid.Token)
                users.append(s.FacultyCode)

        if "S" in sendto:
            staffs = Staff.objects.filter(Deleted = False, Active = True)
            for s in staffs:
                sn = NoticeUser()
                sn.NoticeId = nid
                sn.UserCode = s.StaffCode
                sn.save()
                mid = MessageId.objects.filter(UserId = s.StaffCode)
                if len(mid) > 0:
                    mid = mid[0]
                    tokens.append(mid.Token)
                users.append(s.StaffCode)

    sender["tokens"] = tokens
    send_messages(sender)

    return JsonResponse({"status":"success", "users":users})

def my_handler(sender, instance, **kwargs):
    print(sender)

@csrf_exempt
def save_calendar_event(request):
    type = request.POST.get("type")
    eventdate = request.POST.get("eventdate")
    repeat = request.POST.get("repeat")
    title = request.POST.get("title")
    compcode = request.session["COMPANY"]

    calendar = Calendar()
    calendar.Type = type
    calendar.Repeat = repeat
    calendar.EventDate = eventdate
    calendar.Title = title
    calendar.CompanyCode = compcode
    calendar.save()

    return JsonResponse({"status":"success"})

@csrf_exempt
def load_calendar_events(request):
    eventdate = request.POST.get("date")
    compcode = request.session["COMPANY"]
    calendar = Calendar.objects.filter(EventDate = eventdate, CompanyCode = compcode)
    calendar = parse_data(calendar)

    return JsonResponse(calendar, safe=False)

@csrf_exempt
def load_month_event(request):
    month = request.POST.get("month")
    print(month)
    data = []
    with connection.cursor() as cursor:
        cursor.execute("""select * from main_Calendar c 
        where strftime('%%Y-%%m', c.eventdate) = %s""",[month])
        data = parse_curser(cursor)

    return JsonResponse(data, safe=False)

@csrf_exempt
def load_chats(request):
    sender = request.POST.get("sender")
    receiver = request.POST.get("receiver")

    message = []
    with connection.cursor() as cursor:
        cursor.execute("""select *,
        (case when r.Sender = %s then 'S' when r.Receiver=%s then 'R' end) as direction, 
        strftime('%%d/%%m/%%Y %%H:%%M:%%S', r.CreatedDate) as sendtime
        from main_FacultyRemark r
        where (r.Sender = %s and r.receiver=%s) or (r.sender = %s and r.receiver=%s)
        order by createddate""", [sender, sender ,sender, receiver, receiver, sender])

        message = parse_curser(cursor)

    #post_save.connect(my_handler, sender=FacultyRemark)
    
    return JsonResponse(message, safe=False)

@csrf_exempt
def save_message_id(request):
    userid = request.session["USERID"]
    toke = request.POST.get("token")

    old = MessageId.objects.filter(UserId = userid)
    old.delete()

    msg = MessageId()
    msg.UserId = userid
    msg.Token = toke
    msg.save()

    return JsonResponse({"status":"success"})

@csrf_exempt
def load_privileges(request):
    userid = request.POST.get("userid")

    priv = UserPrivilege.objects.filter(UserId = userid)
    priv = parse_data(priv)

    return JsonResponse(priv, safe=False)

@csrf_exempt
def load_role_privileges(request):
    rolecode = request.POST.get("rolecode")

    priv = RolePrivilege.objects.filter(RoleCode = rolecode)
    priv = parse_data(priv)

    return JsonResponse(priv, safe=False)

@csrf_exempt
def save_permissions(request):
    userid = request.POST.get("userid")
    codes = request.POST.get("codes").split("^")

    print(codes)
    old = UserPrivilege.objects.filter(UserId = userid)
    old.delete()

    for code in codes:
        p = UserPrivilege()
        p.UserId = userid
        p.PrivilegeCode = code
        p.save()

    return JsonResponse({"status":"success"})

@csrf_exempt
def save_role_permissions(request):
    rolecode = request.POST.get("rolecode")
    codes = request.POST.get("codes").split("^")

    old = RolePrivilege.objects.filter(RoleCode = rolecode)
    old.delete()

    for code in codes:
        p = RolePrivilege()
        p.RoleCode = rolecode
        p.PrivilegeCode = code
        p.save()

    return JsonResponse({"status":"success"})

@csrf_exempt
def get_user_detail(request):
    userid = request.POST.get("userid")

    usr = MyUser.objects.filter(UserId=userid)
    usr = parse_data(usr)

    return JsonResponse(usr, safe=False)

@csrf_exempt
def update_member_user(request):
    userid = request.POST.get("userid")
    pas = request.POST.get("password")

    usr = MyUser.objects.filter(UserId = userid)
    newuser = False
    if len(usr) > 0:
        usr = usr[0]
        newuser = False
    else:
        usr = MyUser()
        newuser = True

    usr.UserId = userid
    usr.Password = pas
    usr.UserType = "M"
    usr.CompanyCode = request.session["COMPANY"]
    usr.save()

    if newuser:
        #Giving Default Permissions
        defs = RolePrivilege.objects.filter(RoleCode = "M")
        for d in defs:
            u = UserPrivilege()
            u.UserId = userid
            u.PrivilegeCode = d.PrivilegeCode
            u.save()

    return JsonResponse({"status": "success"})

@csrf_exempt
def save_user(request):
    userid = request.POST.get("userid")
    pas = request.POST.get("password")
    type = request.POST.get("type")
    comp = request.POST.get("company")

    usr = MyUser.objects.filter(UserId = userid)
    newuser = False
    if len(usr) > 0:
        usr = usr[0]
        newuser = False
    else:
        usr = MyUser()
        newuser = True

    usr.UserId = userid
    usr.Password = pas
    usr.UserType = type
    usr.CompanyCode = comp
    #usr.EmployeeCode = emp
    usr.save()

    if newuser:
        #Giving Default Permissions
        defs = RolePrivilege.objects.filter(RoleCode = type)
        for d in defs:
            u = UserPrivilege()
            u.UserId = userid
            u.PrivilegeCode = d.PrivilegeCode
            u.save()

    return JsonResponse({"status": "success"})

@csrf_exempt
def get_parent_payment(request):
    studentcode = request.session["USERID"]
    sessionid = request.POST.get("sessionid")

    payments = []
    with connection.cursor() as cursor:
        cursor.execute("""select i.*, f.documentdate as documentdate, t.Title as typetitle from main_Fee f
        left join main_FeeItem i on i.FeeId= f.id
        left join main_FeeType t on t.id = i.type
        where f.sessionid=%s and f.studentcode=%s""", [sessionid, studentcode])

        payments = parse_curser(cursor)

        print(payments)

    return JsonResponse(payments, safe=False)

@csrf_exempt
def get_parent_homework(request):
    studentcode = request.session["USERID"]
    date = request.POST.get("date")
    homework = []
    with connection.cursor() as cursor:
        cursor.execute("""select h.*, b.Title as subjecttitle 
        from main_Student s
        inner join main_HomeWork h on h.ClassId = s.Class and h.SectionId = s.Section
        inner join main_Subject b on b.id = h.SubjectId
        where s.StudentCode = %s and h.GivenDate = %s""", [studentcode, date])

        homework = parse_curser(cursor)
        print(homework)

    return JsonResponse(homework, safe=False)

@csrf_exempt
def get_parent_report_card(request):
    examid = request.POST.get("examid")
    studentcode = request.session["USERID"]

    marks = []
    with connection.cursor() as cursor:
        cursor.execute("""select r.*, s.Title as SubjectTitle, p.TotalMarks 
        from main_ReportCard r
        inner join main_ExamPaper p on p.id = r.PaperId
        inner join main_Subject s on s.id = p.SubjectId
        where r.ExamId=%s and r.StudentCode=%s""",[examid, studentcode])

        marks = parse_curser(cursor)

    return JsonResponse(marks, safe=False)

@csrf_exempt
def get_exam_by_session(request):
    sessionid = request.POST.get("sessionid")
    exam = Exam.objects.filter(SessionId = sessionid)
    exam = parse_data(exam)

    return JsonResponse(exam, safe=False)

@csrf_exempt
def get_class_section(request):
    cid = request.POST.get("classid")
    section = Section.objects.filter(ClassId = cid)
    section = parse_data(section)

    return JsonResponse(section, safe=False)

@csrf_exempt
def load_homework(request):
    sessionid = request.POST.get("sessionid")
    classid = request.POST.get("classid")
    sectionid = request.POST.get("sectionid")
    subjectid = request.POST.get("subjectid")
    date = request.POST.get("date")

    homework = HomeWork.objects.filter(SessionId = sessionid, ClassId = classid, SectionId = sectionid, GivenDate = date, SubjectId = subjectid)
    homework = parse_data(homework)

    return JsonResponse(homework, safe=False)

@csrf_exempt
def save_homework(request):
    sessionid = request.POST.get("sessionid")
    classid = request.POST.get("classid")
    sectionid = request.POST.get("sectionid")
    date = request.POST.get("date")
    subject = request.POST.get("subjectid")
    text = request.POST.get("text")
    givenby = request.POST.get("givenby")

    homework = HomeWork.objects.filter(SessionId = sessionid, ClassId = classid, SectionId = sectionid, GivenDate = date, SubjectId = subject)
    if len(homework) > 0:
        homework = homework[0]
    else:
        homework = HomeWork()

    homework.SessionId = sessionid
    homework.ClassId = classid
    homework.SectionId = sectionid
    homework.GivenDate = date
    homework.Work = text
    homework.SubjectId = subject
    homework.GivenBy = givenby
    homework.save()

    return JsonResponse({"status":"success"})

@csrf_exempt
def get_parent_syllabus(request):
    request.session["PAGECODE"] = "PG_00026"
    studentcode = request.session["USERID"]
    ses = SchoolSession.objects.filter(Deleted = False, Active=True)
    activesessionid = ses[0].id
    
    student = []
    subjects = []
    syllabus = []

    with connection.cursor() as cursor:
        cursor.execute("""select s.*, c.Title as ClassTitle, t.Title as sectiontitle 
        from main_Student s
        left join main_Class c on c.id = s.Class
        left join main_Section t on t.id = s.Section
        where s.StudentCode=%s""", [studentcode])

        student = parse_curser(cursor)

    with connection.cursor() as cursor:
        cursor.execute("""select b.*, 
        (select COUNT(ss.id) from main_Syllabus ss where s.Class = ss.ClassId and ss.SessionId = %s and ss.SubjectId = b.id and 
        (select COUNT(sb.id) from main_Syllabus sb where sb.parentid = ss.id) = 0) as totalchapter,
        (select COUNT(ss.id) from main_Syllabus ss where s.Class = ss.ClassId and ss.SessionId = %s and ss.Completed = 1 and ss.SubjectId = b.id and
        (select COUNT(sb.id) from main_Syllabus sb where sb.parentid = ss.id) = 0) as donechapter
        from main_Student s
        inner join main_Routine r on r.ClassId = s.Class
        inner join main_Subject b on b.id = r.SubjectId
        where s.StudentCode=%s""", [activesessionid, activesessionid,studentcode])

        subjects = parse_curser(cursor)

    for sub in subjects:
        sdata = {}
        with connection.cursor() as cursor:
            cursor.execute("""select b.* from main_Student s
            inner join main_Syllabus b on b.ClassId = s.Class
            where s.studentcode=%s and b.SessionId=%s and b.SubjectId=%s
            order by b.id""", [studentcode, activesessionid, sub["id"]])

            sdata["title"] = sub["title"]
            sdata["totalchapter"] = sub["totalchapter"]
            sdata["donechapter"] = sub["donechapter"]
            sdata["data"] = parse_curser(cursor)

            syllabus.append(sdata)

    return JsonResponse({"student":student, "syllabus":syllabus})

@csrf_exempt
def get_parent_attendance(request):
    studentcode = request.session["USERID"]
    month = request.POST.get("month")

    attendance = []
    with connection.cursor() as cursor:
        cursor.execute("""select * from main_Attendance a
        where a.StudentCode=%s and strftime('%%Y-%%m',a.attendancedate)=%s""", [studentcode, month])

        attendance = parse_curser(cursor)

    return JsonResponse(attendance, safe=False)

@csrf_exempt
def save_class(request):
    cls = None
    if "classid" in request.POST:
        cls = Class.objects.filter(id = request.POST.get("classid"))
        cls = cls[0]
    else:
        cls = Class()
    cls.ClassTeacher = request.POST.get("faculty")
    cls.Title = request.POST.get("title")
    cls.HasSections = request.POST.get("hassection")
    cls.ClassFee = request.POST.get("fee")
    cls.RoomNo = request.POST.get("roomno")
    cls.save()

    return JsonResponse({"classid":str(cls.id)})

@csrf_exempt
def save_subject(request):
    cls = None
    if "subjectid" in request.POST:
        cls = Subject.objects.filter(id = request.POST.get("subjectid"))
        cls = cls[0]
    else:
        cls = Subject()

    cls.Title = request.POST.get("title")
    cls.save()

    return JsonResponse({"subjectid":str(cls.id)})

@csrf_exempt
def save_exam(request):
    exam = None
    if "examid" in request.POST:
        exam = Exam.objects.filter(id = request.POST.get("examid"))
        exam = exam[0]
    else:
        exam = Exam()

    exam.Title = request.POST.get("title")
    exam.SessionId= request.POST.get("sessionid")
    exam.StartDate = request.POST.get("startdate")
    exam.EndDate = request.POST.get("enddate")
    exam.Remarks = request.POST.get("remarks")
    exam.CreatedBy = "admin"
    exam.save()

    return JsonResponse({"examid":str(exam.id)})

@csrf_exempt
def get_parent_faculty(request):
    userid = request.session["USERID"]

    users = []
    with connection.cursor() as cursor:
        cursor.execute("""select DISTINCT f.FacultyCode, f.FullName, mu.UserId, f.Photo as UserPhoto,
        (select GROUP_CONCAT(b.Title, ' + ') from main_Subject b where b.id = r.SubjectId) as AllSubjects 
        from main_Student s
        inner join main_Routine r on r.ClassId = s.Class and (r.SectionId = '' or r.SectionId = s.Section)
        inner join main_Faculty f on f.FacultyCode = r.FacultyId
        inner join main_MyUser mu on mu.EmployeeCode = f.FacultyCode or mu.userid = f.facultycode
        where s.StudentCode=%s""", [userid])

        users = parse_curser(cursor)
        print(users)
    
    return JsonResponse(users, safe=False)

@csrf_exempt
def get_discus_parent(request):
    userid = request.session["USERID"]

    users = []
    with connection.cursor() as cursor:
        cursor.execute("""select Distinct s.fathername as fullname, 
        s.Fullname ||' : Class - ' || c.Title as allsubjects, s.StudentCOde as userid,
        s.Photo as UserPhoto
        from main_MyUser mu
        inner join main_Faculty f on f.FacultyCode = mu.EmployeeCode
        inner join main_Routine r on r.FacultyId = f.FacultyCode
        inner join main_Class c on c.id = r.ClassId
        inner join main_Student s on s.Class = c.id
        where mu.userid = %s group by s.StudentCode""", [userid])

        users = parse_curser(cursor)

    return JsonResponse(users, safe=False)

@csrf_exempt
def save_paper(request):
    paper = None
    if "paperid" in request.POST:
        paper = ExamPaper.objects.filter(id = request.POST.get("paperid"))
        paper = paper[0]
    else:
        paper = ExamPaper()

    paper.SubjectId = request.POST.get("subject")
    paper.ClassId = request.POST.get("class")
    paper.TotalMarks = request.POST.get("totalmarks")
    paper.SectionId = request.POST.get("section")
    if "paperdate" in request.POST:
        paper.PaperDate = request.POST.get("paperdate")
    else:
        paper.PaperDate = None
    paper.CreatedBy = "admin"
    paper.ExamId = request.POST.get("exam")
    paper.Content = request.POST.get("content")
    paper.CreatedBy = "admin"
    paper.save()

    return JsonResponse({"paperid":str(paper.id)})

@csrf_exempt
def get_exam_list(request):
    sessionid = request.POST.get("sessionid")
    alldata = None
    with connection.cursor() as cursor:
        cursor.execute("""select *, (case when e.enddate < CURRENT_DATE then '1' else '0' end) as Finished
        from main_Exam e
        where e.Deleted = 0 and e.SessionId = %s order by startdate desc""", [sessionid])
        alldata = parse_curser(cursor)
    return JsonResponse(alldata, safe=False)

@csrf_exempt
def print_report_card(request):
    template = request.FILES["template"]
    doc = Document(template)

    headertable = doc.tables[1]
    marktable = doc.tables[0]

    keys = request.POST.keys()    
    
    datastring = request.POST.get("DATASTRING_L")
    items = datastring.split("#")

    for row in headertable.rows:
        for cell in row.cells:
            for key in keys:
                if cell.text == "<<"+ key.replace("_F","") +">>":
                    cell.text = request.POST.get(key)
                    font = cell.paragraphs[0].runs[0].font
                    font.size = Pt(11)
                    font.name = 'Century Gothic'

    for item in items[:-1]:
        itemsplit = item.split("^")
        new_row = marktable.add_row()
        new_row.cells[0].text = itemsplit[0]
        new_row.cells[1].text = itemsplit[1]
        new_row.cells[2].text = itemsplit[2]
        
        for i in range(3):
            font = new_row.cells[i].paragraphs[0].runs[0].font
            font.size = Pt(11)
            font.name = 'Century Gothic'

            cell = new_row.cells[i]
            # Access the paragraph in the cell
            paragraph = cell.paragraphs[0]
            if i > 0:
            # Set the alignment of the paragraph to center
                paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

            # Set the vertical alignment of text in the cell to center
            tc = cell._element.tcPr
            vAlign = docx.oxml.shared.OxmlElement('w:vAlign')
            vAlign.set(docx.oxml.ns.qn('w:val'), 'center')
            tc.append(vAlign)

            new_row.height = docx.shared.Cm(1)

            if items.index(item) == len(items)-2:
                for run in paragraph.runs:
                    run.bold = True


    newname = generate_id(15)+".docx"
    doc.save(newname)

    convert(newname, "static/Doc/"+newname.replace(".docx", ".pdf"))

    with open("static/Doc/"+newname.replace(".docx", ".pdf"), 'rb') as file:
            file_content = file.read()

    # Encode the file content as Base64
    file_content_base64 = base64.b64encode(file_content).decode('utf-8')

    # Create the response JSON object
    response_data = {
        'filename': 'pdffile.pdf',
        'content': file_content_base64
    }
        
    os.remove(newname)
    os.remove("static/Doc/"+newname.replace(".docx", ".pdf"))
    return JsonResponse(response_data, status=201)


@csrf_exempt
def print_receipt(request):
    template = request.FILES["template"]
    doc = Document(template)

    headertable = doc.tables[0]
    itemtable = doc.tables[1]
    totaltable = doc.tables[2]
    wordtable = doc.tables[3]

    keys = request.POST.keys()
        
    for row in headertable.rows:
        for cell in row.cells:
            for key in keys:
                if cell.text == "<<"+ key.replace("_F","") +">>":
                    cell.text = request.POST.get(key)
                    font = cell.paragraphs[0].runs[0].font
                    font.size = Pt(11)
                    font.name = 'Calibri'

    for row in totaltable.rows:
        for cell in row.cells:
            if cell.text == "<<FINAL>>":
                cell.text = request.POST.get("FINALAMOUNT")
                font = cell.paragraphs[0].runs[0].font
                font.size = Pt(11)
                font.name = 'Calibri'

    for row in wordtable.rows:
        for cell in row.cells:
            if cell.text == "<<WORDS>>":
                cell.text = request.POST.get("WORDS")
                font = cell.paragraphs[0].runs[0].font
                font.size = Pt(11)
                font.name = 'Calibri'

    datastring = request.POST.get("DATASTRING_L")
    items = datastring.split("#")
    for item in items[:-1]:
        itemsplit = item.split("^")
        new_row = itemtable.add_row()
        new_row.cells[0].text = itemsplit[0]
        new_row.cells[1].text = itemsplit[1]
        new_row.cells[2].text = itemsplit[2]
        new_row.cells[3].text = itemsplit[3]
        new_row.cells[4].text = itemsplit[4]
        
        for i in range(5):
            font = new_row.cells[i].paragraphs[0].runs[0].font
            font.size = Pt(11)
            font.name = 'Calibri'

    newname = generate_id(15)+".docx"
    doc.save(newname)

    convert(newname, "static/Doc/"+newname.replace(".docx", ".pdf"))

    with open("static/Doc/"+newname.replace(".docx", ".pdf"), 'rb') as file:
            file_content = file.read()

    # Encode the file content as Base64
    file_content_base64 = base64.b64encode(file_content).decode('utf-8')

    # Create the response JSON object
    response_data = {
        'filename': 'pdffile.pdf',
        'content': file_content_base64
    }
        
    os.remove(newname)
    os.remove("static/Doc/"+newname.replace(".docx", ".pdf"))
    return JsonResponse(response_data, status=201)

@csrf_exempt
def save_routine(request):
    rout = None
    if "routineid" in request.POST:
        rout = Routine.objects.filter(id = request.POST.get("routineid"))
        rout = rout[0]
    else:
        rout = Routine()

    rout.SessionId = request.POST.get("sessionid")
    rout.ClassId = request.POST.get("class")
    rout.SectionId = request.POST.get("section")
    rout.SubjectId = request.POST.get("subject")
    rout.FacultyId = request.POST.get("faculty")
    rout.Day = request.POST.get("day")
    rout.StartTime = request.POST.get("starttime")
    rout.EndTime = request.POST.get("endtime")
    rout.CreatedBy = "admin"
    rout.save()

    return JsonResponse({"status":"success"})

@csrf_exempt
def save_section(request):
    cls = None
    if "sectionid" in request.POST:
        cls = Section.objects.filter(id = request.POST.get("sectionid"))
        cls = cls[0]
    else:
        cls = Section()

    cls.ClassTeacher = request.POST.get("faculty")
    cls.Title = request.POST.get("title")
    cls.ClassId = request.POST.get("classid")
    cls.RoomNo = request.POST.get("roomno")
    cls.save()

    return JsonResponse({"sectionid":str(cls.id)})

@csrf_exempt
def get_faculty_for_salary(request):
    sessionid = request.POST.get("sessionid")
    monthyear = request.POST.get("monthyear").split("-")
    type = request.POST.get("type")
    month = monthyear[1]
    year = monthyear[0]
    
    alldata = []
    if type == "F":
        with connection.cursor() as cursor:
            cursor.execute("""select f.*,
            (select COUNT(fa.id) from main_FacultyAttendance fa 
            where strftime('%%m',fa.AttendanceDate) = %s and strftime('%%Y',fa.AttendanceDate) = %s and 
            f.FacultyCode = fa.FacultyCode and fa.sessionid = %s and fa.Status = 'A') as absentcount,
            (select sr.MonthYear from main_SalaryRecord sr where sr.FacultyCode = f.FacultyCode
            order by substr(sr.monthyear,1, 4)||substr(sr.monthyear,6, 7) desc LIMIT 1) as lastrecord
            from main_Faculty f""", [month, year, sessionid])

            alldata = parse_curser(cursor)

    elif type == "S":
        with connection.cursor() as cursor:
            cursor.execute("""select f.*, f.StaffCode as FacultyCode,
            (select COUNT(fa.id) from main_FacultyAttendance fa 
            where strftime('%%m',fa.AttendanceDate) = %s and strftime('%%Y',fa.AttendanceDate) = %s and 
            f.StaffCode = fa.FacultyCode and fa.sessionid = %s and fa.Status = 'A') as absentcount,
            (select sr.MonthYear from main_SalaryRecord sr where sr.FacultyCode = f.StaffCode
            order by substr(sr.monthyear,1, 4)||substr(sr.monthyear,6, 7) desc LIMIT 1) as lastrecord
            from main_Staff f""", [month, year, sessionid])

            alldata = parse_curser(cursor)
        

    return JsonResponse(alldata, safe=False)

@csrf_exempt
def get_class_subject(request):
    cls = request.POST.get("classid")
    sectionid = request.POST.get("sectionid")
    sessionid = request.POST.get("sessionid")
    examid = request.POST.get("examid")

    student = None
    subjects = None

    with connection.cursor() as cursor:
        cursor.execute('''select DISTINCT r.id as paperid, s.Title, r.TotalMarks, ss.Title as SessionTitle
        from main_ExamPaper r
        inner join main_Exam ex on ex.id = r.ExamId
        inner join main_Subject s on s.id = r.SubjectId
        inner join main_SchoolSession ss on ss.id = ex.SessionId
        where r.ClassId=%s and (r.SectionId='' or r.SectionId is null or r.SectionId like %s) 
        and ex.SessionId=%s and ex.id=%s''', [cls, '%'+sectionid+'%', sessionid, examid])

        subjects = parse_curser(cursor)

    with connection.cursor() as cursor:
        cursor.execute("""select s.FullName, s.StudentCode
        from main_Student s
        where s.Deleted = 0 and s.Class=%s and (s.section = '' or s.section is null or s.Section = %s)""", [cls, sectionid])
        student = parse_curser(cursor)

    return JsonResponse({"students":student, "subs": subjects})

@csrf_exempt
def get_syllabus(request):
    classid = request.POST.get("classid")
    #sectionid = request.POST.get("sectionid")
    subjectid = request.POST.get("subjectid")
    #sessionid = request.POST.get("sessionid")

    syllabus = Syllabus.objects.filter(ClassId = classid, SubjectId = subjectid).order_by("ChapterNo")

    syllabus = parse_data(syllabus)

    return JsonResponse(syllabus, safe=False)

@csrf_exempt
def mark_individual_attendance(request):
    date = request.POST.get("date")
    sessionid = request.POST.get("sessionid")
    code = request.POST.get("studentcode")
    status = request.POST.get("status")

    atype = "S"
    if "type" in request.POST:
        atype = request.POST.get("type")

    old = Attendance.objects.filter(StudentCode = code, AttendanceDate = date, SessionId = sessionid)
    old.delete()

    att = Attendance()
    att.StudentCode = code
    att.AttendanceDate = date
    att.SessionId = sessionid
    att.Status = status
    att.Type = atype
    att.MarkedBy = request.session["USERID"]
    att.save()

    return JsonResponse({"status":"success"})

@csrf_exempt
def mark_bulk_attendance(request):
    date = request.POST.get("date")
    sessionid = request.POST.get("sessionid")
    datastring = request.POST.get("datastring").split("#")

    atype = "S"
    if "type" in request.POST:
        atype = request.POST.get("type")

    for data in datastring[:-1]:
        spl = data.split("^")
        old = Attendance.objects.filter(StudentCode = spl[0], AttendanceDate = date, SessionId = sessionid)
        old.delete()
        att = Attendance()
        att.StudentCode = spl[0]
        att.AttendanceDate = date
        att.SessionId = sessionid
        att.Status = spl[1]
        att.Type = atype
        att.MarkedBy = request.session["USERID"]
        att.save()

    return JsonResponse({"status":"success"})

@csrf_exempt
def mark_faculty_attendance(request):
    date = request.POST.get("date")
    code = request.POST.get("facultycode")
    status = request.POST.get("status")
    sessionid = request.POST.get("sessionid")

    old = FacultyAttendance.objects.filter(FacultyCode = code, AttendanceDate = date)
    old.delete()

    att = FacultyAttendance()
    att.FacultyCode = code
    att.AttendanceDate = date
    att.Status = status
    att.SessionId = sessionid
    att.MarkedBy = request.session["USERID"]
    att.save()

    return JsonResponse({"status":"success"})

@csrf_exempt
def mark_self_attendance(request):
    code = request.POST.get("facultycode")
    lati = request.POST.get("latitude")
    longi = request.POST.get("longitude")
    session = request.POST.get("sessionid")
    status = request.POST.get("status")
    
    currenttime = str(timezone.now())
    cdate = currenttime[0:10]

    with connection.cursor() as cursor:
        cursor.execute("""delete from main_FacultyAttendance 
        where strftime('%%Y-%%m-%%d', AttendanceDate)=%s and FacultyCode = %s""", [cdate, code])

    att = FacultyAttendance()
    att.FacultyCode = code
    att.Latitude = lati
    att.Longitude = longi
    att.SessionId = session
    att.MarkedBy = "admin"
    att.Status = status
    att.AttendanceDate = currenttime
    att.save()

    return JsonResponse({"status":"success"})

@csrf_exempt
def mark_member_self_attendance(request):
    code = request.POST.get("membercode")
    lati = request.POST.get("latitude")
    longi = request.POST.get("longitude")
    status = request.POST.get("status")
    
    currenttime = str(timezone.now())[0:19]
    cdate = currenttime[0:10]

    with connection.cursor() as cursor:
        cursor.execute("""delete from main_Attendance 
        where strftime('%%Y-%%m-%%d', AttendanceDate)=%s and StudentCode = %s""", [cdate, code])

    att = Attendance()
    att.StudentCode = code
    att.Latitude = lati
    att.Longitude = longi
    att.MarkedBy = request.session["USERID"]
    att.Status = status
    att.AttendanceDate = currenttime
    att.save()

    return JsonResponse({"status":"success"})

@csrf_exempt
def get_main_dashboard(request):
    session = request.POST.get("session")

    currentmonth = str(timezone.now())[0:7]

    stats = []  
    with connection.cursor() as cursor:
        cursor.execute("""select 
        coalesce((select SUM(f.TotalAmount) from main_Fee f 
        where f.SessionId = %s and strftime('%%Y-%%m',f.DocumentDate)=%s),0) as allcollection,
        coalesce((select SUM(e.Amount) from main_Expense e 
        where e.SessionId = %s and strftime('%%Y-%%m',e.DocumentDate)=%s),0) as allexpense""", [session, currentmonth, session, currentmonth])
        stats = parse_curser(cursor)

    syllabus = []
    with connection.cursor() as cursor:
        cursor.execute("""select Distinct c.*,
        (select COUNT(s.id) from main_Syllabus s where c.id = s.ClassId and s.SessionId = %s and
        (select COUNT(sb.id) from main_Syllabus sb where sb.parentid = s.id) = 0) as totalchapter,
        (select COUNT(s.id) from main_Syllabus s where c.id = s.ClassId and s.SessionId = %s and s.Completed = 1 and
        (select COUNT(sb.id) from main_Syllabus sb where sb.parentid = s.id) = 0) as donechapter
        from main_Class c 
        Group By c.id""", [session, session])
        
        syllabus = parse_curser(cursor)

    overdue = []
    with connection.cursor() as cursor:
        cursor.execute("""select SUM(c.ClassFee) as overdue from main_Student s
        inner join main_Class c on c.id = s.Class
        where not exists(select f.id from main_Fee f 
        inner join main_FeeItem i on i.Feeid = f.id
        where f.StudentCode = s.StudentCode and i.MonthYear=%s and f.sessionid = %s)""", [currentmonth, session])
        overdue = parse_curser(cursor)

    return JsonResponse({"syllabus":syllabus, "data": stats, "overdue":overdue})

@csrf_exempt
def get_salary_history(request):
    code = request.POST.get("facultycode")

    salary = SalaryRecord.objects.filter(FacultyCode = code)

    salary = parse_data(salary)
    return JsonResponse(salary, safe=False)

@csrf_exempt
def get_member_dashboard(request):
    code = request.POST.get("membercode")
    
    alldata = []
    payments = []
    with connection.cursor() as cursor:
        cursor.execute("""select m.*,
        (select fa.Status from main_Attendance fa 
        where strftime('%%Y-%%m-%%d', AttendanceDate)=%s order by AttendanceDate desc limit 1) as AttendanceStatus
        from main_LibraryMember m
        where m.MemberCode = %s""", [str(timezone.now())[0:10], code])

        alldata = parse_curser(cursor)

    with connection.cursor() as cursor:
        cursor.execute("""select f.DocumentDate,
        fi.*
        from main_Fee f
        left join main_FeeItem fi on fi.Feeid = f.id
        where f.StudentCode=%s""", [code])

        payments = parse_curser(cursor)

    return JsonResponse({"data":alldata, "payments":payments})

@csrf_exempt
def get_faculty_dashboard(request):
    code = request.POST.get("code")
    session = request.POST.get("session")
    cdate = request.POST.get("date")
    day = request.POST.get("day")

    alldata = []
    with connection.cursor() as cursor:
        cursor.execute("""select f.*, 
        (select fa.Status from main_FacultyAttendance fa 
        where strftime('%%Y-%%m-%%d', AttendanceDate)=%s and fa.SessionId = %s order by AttendanceDate desc LIMIT 1) as AttendanceStatus,
        (select COUNT(r.id) from main_Routine r 
        where r.facultyid = %s and r.SessionId = %s and r.day = %s) as classes
        from main_Faculty f
        where f.FacultyCode=%s""", [cdate, session, code, session, day, code])

        alldata = parse_curser(cursor)
    
    syllabus = []
    with connection.cursor() as cursor:
        cursor.execute("""select Distinct r.*,
        (select COUNT(s.id) from main_Syllabus s 
        where s.SubjectId = r.SubjectId and s.ClassId = r.ClassId and s.sessionid = %s and
        (select COUNT(sb.id) from main_Syllabus sb where sb.parentid = s.id) = 0) as totalsyllabus,
        (select COUNT(s.id) from main_Syllabus s 
        where s.SubjectId = r.SubjectId and s.ClassId = r.ClassId and s.Completed = 1 and s.sessionid = %s and
        (select COUNT(sb.id) from main_Syllabus sb where sb.parentid = s.id) = 0) as completedsyllabus,
        sb.Title as SubjectTitle,
        cl.Title as ClassTitle 
        from main_Routine r
        inner join main_Subject sb on sb.id = r.SubjectId
        inner join main_Class cl on cl.id = r.ClassId
        where r.FacultyId = %s and r.SessionId = %s Group By r.SubjectId""", [session, session, code, session])
        
        syllabus = parse_curser(cursor)
    return JsonResponse({"data":alldata, "syllabus":syllabus})

@csrf_exempt
def get_principal(request):
    principal = []
    with connection.cursor() as cursor:
        cursor.execute("""select * from main_MyUser u
                       where u.UserType = 'L' and u.Blocked = 0 and u.Deleted = 0""")
        principal = parse_curser(cursor)
    return JsonResponse(principal, safe=False)

@csrf_exempt
def get_student_attendance(request):
    cls = request.POST.get("classid")
    date = request.POST.get("date")
    sec = request.POST.get("sectionid")

    alldata = []
    with connection.cursor() as cursor:
        cursor.execute("""select s.*,
        (select a.Status from main_Attendance a where a.StudentCode = s.StudentCode and
                       strftime('%%Y-%%m-%%d', a.AttendanceDate) = %s) as attendancestatus 
        from main_Student s where s.Class = %s and (s.Section = %s or %s = '' or %s is null)""", [date, cls, sec, sec, sec])

        alldata = parse_curser(cursor)

    return JsonResponse(alldata, safe=False)

@csrf_exempt
def delete_chapter(request):
    sortid = request.POST.get("sortid")
    item = Syllabus.objects.filter(id = sortid)
    item.delete()

    return JsonResponse({"status":"success"})

@csrf_exempt
def mark_chapter_done(request):
    sortid = request.POST.get("sortid")
    status = request.POST.get("status")

    item = Syllabus.objects.filter(id = sortid)
    if len(item) > 0:
        if status == "1":
            item[0].Completed = True
            item[0].CompletedOn = timezone.now()
        else:
            item[0].Completed = False
        item[0].save()

    return JsonResponse({"status":"success"})

@csrf_exempt
def save_salary(request):
    sessionid = request.POST.get("session")
    monthyear = request.POST.get("monthyear")
    datastring = request.POST.get("datastring")

    datastring = datastring.split("#")
    ids = []
    exist = []
    for data in datastring[:-1]:
        datasplit = data.split("^")
        salary = SalaryRecord()
        old = SalaryRecord.objects.filter(FacultyCode = datasplit[0], MonthYear = monthyear)
        if len(old) > 0:
            exist.append(datasplit[0])
            break
        salary.FacultyCode = datasplit[0]
        salary.Salary = datasplit[1]
        salary.Days = datasplit[2]
        salary.Payable = datasplit[3]
        salary.Deduction = datasplit[4]
        salary.Incentive = datasplit[5]
        salary.NetSalary = datasplit[6]
        salary.SessionId = sessionid
        salary.MonthYear = monthyear
        salary.CreatedBy = "admin"
        salary.save()
        ids.append(salary.id)

    if len(exist) > 0:
        for id in ids:
            old = SalaryRecord.objects.filter(id = id)
            old.delete()
        return JsonResponse({"status":"exists", "data":exist})
    else:
        return JsonResponse({"status":"success", "data":""})

@csrf_exempt
def save_chapter(request):
    classid = request.POST.get("classid")
    sectionid = request.POST.get("sectionid")
    subjectid = request.POST.get("subjectid")
    sessionid = request.POST.get("sessionid")

    parentid = request.POST.get("parentid")
    title = request.POST.get("title")
    chapterno = request.POST.get("chapterno")
    sortid = request.POST.get("sortid")

    syllabus = Syllabus()

    if sortid != "":
        syllabus = Syllabus.objects.filter(id = sortid)
        if len(syllabus) > 0:
            syllabus = syllabus[0]
        else:
            syllabus = Syllabus()
    syllabus.ClassId = classid
    syllabus.SectionId = sectionid
    syllabus.SessionId = sessionid
    syllabus.SubjectId = subjectid

    syllabus.ParentId = parentid
    syllabus.ChapterNo = chapterno
    syllabus.Title = title
    syllabus.save()

    newid = syllabus.id

    return JsonResponse({"id":newid})

@csrf_exempt
def save_syllabus(request):
    classid = request.POST.get("classid")
    sectionid = request.POST.get("sectionid")
    subjectid = request.POST.get("subjectid")
    sessionid = request.POST.get("sessionid")
    datastring = request.POST.get("datastring")
    
    datastring = datastring.split("#")
    count = 1

    old = Syllabus.objects.filter(ClassId = classid, SessionId = sessionid, SubjectId = subjectid)
    old.delete()

    for data in datastring[:-1]:
        spl = data.split("^")
        parentid = spl[0]
        itemid = spl[1]

        psyllabus = Syllabus.objects.get(id = parentid)
        if psyllabus is not None:
            parentid = psyllabus.id
        else:
            psyllabus = Syllabus()
            psyllabus.ClassId = classid
            psyllabus.SectionId = sectionid
            psyllabus.SessionId = sessionid
            psyllabus.SubjectId = subjectid

            psyllabus.ParentId = parentid
            psyllabus.ChapterNo = count
            psyllabus.save()

        syllabus = Syllabus()
        syllabus.ClassId = classid
        syllabus.SectionId = sectionid
        syllabus.SessionId = sessionid
        syllabus.SubjectId = subjectid

        syllabus.ParentId = parentid
        syllabus.ChapterNo = count
        syllabus.save()

        count+=1

    return JsonResponse({"status":"success"})

@csrf_exempt
def get_student_marks(request):
    studentcode = request.POST.get("studentcode")
    papers = request.POST.get("papers").split("^")
    alldata = {}

    for p in papers[:-1]:
        res = ReportCard.objects.filter(PaperId = p, StudentCode = studentcode)
        if len(res) > 0:
            alldata[p] = res[0].Marks
    return JsonResponse(alldata, safe=False)

@csrf_exempt
def get_student_by_class(request):
    classid = request.POST.get("classid")
    sectionid = request.POST.get("sectionid")
    sessionid = request.POST.get("sessionid")

    alldata = None
    with connection.cursor() as cursor:
        cursor.execute('''select s.StudentCode, s.FullName, s.id,
        (select GROUP_CONCAT(a.AttendanceDate|| '#' || a.Status, '^') from main_Attendance a 
        where a.StudentCode = s.StudentCode and a.SessionId = %s) as attendancedata 
        from main_Student s
        where Class=%s and (section = '' or section is null or Section=%s)''', [sessionid ,classid, sectionid])

        alldata = parse_curser(cursor)

    return JsonResponse(alldata, safe=False)

@csrf_exempt
def get_member_by_batch(request):
    batchid = request.POST.get("batchid")
    comp = request.session["COMPANY"]

    alldata = None
    with connection.cursor() as cursor:
        cursor.execute('''select s.MemberCode, s.FullName, s.id,
        (select GROUP_CONCAT(a.AttendanceDate|| '#' || a.Status, '^') from main_Attendance a 
        where a.StudentCode = s.MemberCode and a.Type = 'M') as attendancedata 
        from main_LibraryMember s
        where (%s = '' or BatchId=%s) and (%s = '' or s.CompanyCode = %s) order by s.FullName''', [batchid,batchid, comp, comp])

        alldata = parse_curser(cursor)

    return JsonResponse(alldata, safe=False)

@csrf_exempt
def get_faculty_attendance(request):
    sessionid = request.POST.get("sessionid")
    month = request.POST.get("month")
    type = request.POST.get("type")

    alldata = None
    if type == "F":
        with connection.cursor() as cursor:
            cursor.execute('''select s.FacultyCode, s.FullName, s.id,
            (select GROUP_CONCAT(a.AttendanceDate|| '#' || a.Status, '^') from main_FacultyAttendance a 
            where a.FacultyCode = s.FacultyCode and a.SessionId = %s and strftime('%%Y-%%m', a.AttendanceDate)=%s) as attendancedata 
            from main_Faculty s''', [sessionid, month])

            alldata = parse_curser(cursor)
    elif type == "S":
        with connection.cursor() as cursor:
            cursor.execute('''select s.StaffCode as FacultyCode, s.FullName, s.id,
            (select GROUP_CONCAT(a.AttendanceDate|| '#' || a.Status, '^') from main_FacultyAttendance a 
            where a.FacultyCode = s.StaffCode and a.SessionId = %s and strftime('%%Y-%%m', a.AttendanceDate)=%s) as attendancedata 
            from main_Staff s''', [sessionid, month])

            alldata = parse_curser(cursor)
        
    return JsonResponse(alldata, safe=False)

@csrf_exempt
def save_attendance(request):
    mon = request.POST.get("month").split("-")
    data = request.POST.get("datastring")
    sessionid = request.POST.get("sessionid")

    spl = data.split("!")

    for s in spl[:-1]:
        sp = s.split("#")
        code = sp[0]
        att = sp[1].split("^")

        #Deleting old data
        old = Attendance.objects.filter(SessionId = sessionid, AttendanceDate__year = mon[0], AttendanceDate__month = mon[1], StudentCode = code)
        old.delete()

        ##saving the data
        ind = 0
        for a in att[:-1]:
            ind+=1
            if a != "O":
                attendance = Attendance()
                attendance.StudentCode = code
                attendance.Status = a
                attendance.Type = "S"
                attendance.SessionId = sessionid
                attendance.SessionId = sessionid
                attendance.AttendanceDate = mon[0]+"-"+mon[1]+"-"+'{:02d}'.format(ind)
                attendance.MarkedBy = "admin"
                attendance.save()

    return JsonResponse({"status": "success"})

@csrf_exempt
def save_member_attendance(request):
    mon = request.POST.get("month").split("-")
    data = request.POST.get("datastring")
    comp = request.session["COMPANY"]
    
    spl = data.split("!")

    for s in spl[:-1]:
        sp = s.split("#")
        code = sp[0]
        att = sp[1].split("^")

        #Deleting old data
        old = Attendance.objects.filter(AttendanceDate__year = mon[0], AttendanceDate__month = mon[1], StudentCode = code)
        old.delete()

        ##saving the data
        ind = 0
        for a in att[:-1]:
            ind+=1
            if a != "O":
                attendance = Attendance()
                attendance.StudentCode = code
                attendance.Status = a
                attendance.Type = 'M'
                attendance.AttendanceDate = mon[0]+"-"+mon[1]+"-"+'{:02d}'.format(ind)
                attendance.MarkedBy = "admin"
                attendance.save()

    return JsonResponse({"status": "success"})

@csrf_exempt
def save_faculty_attendance(request):
    mon = request.POST.get("month").split("-")
    data = request.POST.get("datastring")
    sessionid = request.POST.get("sessionid")
    
    spl = data.split("!")

    for s in spl[:-1]:
        sp = s.split("#")
        code = sp[0]
        att = sp[1].split("^")

        #Deleting old data
        old = FacultyAttendance.objects.filter(SessionId = sessionid, AttendanceDate__year = mon[0], AttendanceDate__month = mon[1], FacultyCode = code)
        old.delete()

        ##saving the data
        ind = 0
        for a in att[:-1]:
            ind+=1
            if a != "O":
                attendance = FacultyAttendance()
                attendance.FacultyCode = code
                attendance.Status = a
                attendance.SessionId = sessionid
                attendance.SessionId = sessionid
                attendance.AttendanceDate = mon[0]+"-"+mon[1]+"-"+'{:02d}'.format(ind)
                attendance.MarkedBy = "admin"
                attendance.save()

    return JsonResponse({"status": "success"})

@csrf_exempt
def get_class_fee(request):
    cls = Class.objects.get(id = request.POST.get("classid"))
    classfee = 0
    if cls is not None:
        classfee = cls.ClassFee
    return JsonResponse(classfee, safe=False)

@csrf_exempt
def get_class_routine(request):
    cls = request.POST.get("class")
    sec = request.POST.get("section")
    day = request.POST.get("day")
    sessionid = request.POST.get("sessionid")

    rout = None
    with connection.cursor() as cursor:
        cursor.execute("""select r.*, c.Title as SubjectTitle, f.FirstName as firstname, f.MiddleName as middlename, f.LastName as lastname
        from main_Routine r
        left join main_Subject c on c.id = r.SubjectId
        left join main_Faculty f on f.FacultyCode = r.FacultyId
        where r.Classid = %s and (r.sectionid = '' or r.sectionid is null or r.sectionid = %s) and r.day = %s and r.SessionId=%s
        order by r.starttime""", [cls, sec, day, sessionid])

        rout = parse_curser(cursor)
    return JsonResponse({"data":rout})

@csrf_exempt
def get_faculty_routine(request):
    faculty = request.POST.get("faculty")
    day = request.POST.get("day")
    sessionid = request.POST.get("sessionid")

    rout = None
    with connection.cursor() as cursor:
        cursor.execute("""select r.*, c.Title as SubjectTitle, f.FullName, cl.Title as classtitle, 
        sec.Title as sectiontitle
        from main_Routine r
        left join main_Subject c on c.id = r.SubjectId
        left join main_Faculty f on f.FacultyCode = r.FacultyId
        left join main_Class cl on cl.id = r.ClassId
        left join main_Section sec on sec.id = r.SectionId
        where r.FacultyId = %s and r.day = %s and r.SessionId=%s
        order by r.starttime""", [faculty, day, sessionid])

        rout = parse_curser(cursor)
    return JsonResponse({"data":rout})


@csrf_exempt
def get_all_class(request):
    alldata = []
    with connection.cursor() as cursor:
        cursor.execute("""select c.*,
        (select COUNT(s.id) from main_SubjectClass s where c.id = s.ClassId and s.SessionId = %s) as totalsubjects
        from main_Class c
        where c.Deleted = 0
        order by c.Title""", [current_session()])

        alldata = parse_curser(cursor)
    return JsonResponse({"data":alldata})

@csrf_exempt
def get_all_section(request):
    cls = Section.objects.filter(Deleted = False)
    alldata = parse_data(cls)
    return JsonResponse({"data":alldata})

@csrf_exempt
def get_all_subject(request):
    cls = Subject.objects.filter(Deleted = False)
    alldata = parse_data(cls)
    return JsonResponse({"data":alldata})

@csrf_exempt
def save_fee(request):
    fee = Fee()
    retid = ""

    if "feeid" in request.POST:
        fee = Fee.objects.filter(id = request.POST.get("feeid"))
        if len(fee) > 0:
            fee = fee[0]
        else:
            fee = Fee()

    feetp = "S"
    if "type" in request.POST:
        feetp = request.POST.get("type")

    fee.CompanyCode = request.session["COMPANY"]
    fee.StudentCode = request.POST.get("student")
    fee.DocumentDate = request.POST.get("collectdate")
    fee.Mode = request.POST.get("mode")
    fee.TotalAmount = request.POST.get("total")
    fee.Remarks = request.POST.get("remarks")
    fee.ReferenceNo = request.POST.get("refno")
    fee.CreatedBy = "admin"
    fee.Type = feetp
    fee.SessionId = request.POST.get("sessionid")
    fee.save()
    newid = fee.id
    fee.DocumentCode = "REC-"+'{:05d}'.format(newid)
    fee.save()
    retid = fee.id

    #deleting old items
    old = FeeItem.objects.filter(FeeId = newid)
    old.delete()

    #saving new items
    datastring = request.POST.get("datastring").split("#")
    for data in datastring[:-1]:
        item = FeeItem()
        spl = data.split("^")
        item.FeeId = newid
        item.Type = spl[0]
        item.MonthYear = spl[1]
        item.Amount = spl[2]
        item.Discount = spl[3]
        item.Fine = spl[4]
        item.TotalAmount = spl[5]
        item.save()

    return JsonResponse({"feeid": retid})

@csrf_exempt
def save_class_subject(request):
    classid = request.POST.get("classid")
    sessionid = request.POST.get("sessionid")
    datastring = request.POST.get("datastring").split("#")

    old = SubjectClass.objects.filter(ClassId = classid, SessionId = sessionid)
    old.delete()
    for data in datastring:
        subclass = SubjectClass()
        subclass.ClassId = classid
        subclass.SessionId = sessionid

        subfac = data.split("^")
        subclass.SubjectId = subfac[0]
        subclass.FacultyId = subfac[1]
        subclass.save()

    return JsonResponse({"status":"success"})

@csrf_exempt
def save_expense(request):
    exp = Expense()
    if "expenseid" in request.POST:
        exp = Expense.objects.filter(ExpenseCode = request.POST.get("expenseid"))
        if len(exp) > 0:
            exp = exp[0]
        else:
            exp = Expense()

    exp.SessionId = request.POST.get("sessionid")
    exp.DocumentDate = request.POST.get("expensedate")
    exp.Amount = request.POST.get("amount")
    exp.DocumentNo = request.POST.get("docno")
    exp.SessionId = request.POST.get("modalsession")
    exp.CreatedBy = "admin"
    exp.Type = request.POST.get("expensetype")
    exp.InvoiceCopy = request.POST.get("invoice")
    exp.Remarks = request.POST.get("remarks")
    exp.Mode = request.POST.get("mode")
    exp.ReferenceNo = request.POST.get("refno")
    exp.save()
    newid = exp.id
    newcode = "EXP-"+'{:05d}'.format(newid)
    exp.ExpenseCode = newcode
    exp.save()

    return JsonResponse({"code": newcode})


@csrf_exempt
def save_faculty(request):
    faculty = Faculty()
    newcode = ""
    if "facultycode" in request.POST:
        faculty = Faculty.objects.filter(FacultyCode = request.POST.get("facultycode"))
        if len(faculty) > 0:
            faculty = faculty[0]
            newcode = faculty.FacultyCode
        else:
            faculty = Faculty()
            newcode = "FAC-"+'{:05d}'.format(len(Faculty.objects.all())+1)
            faculty.FacultyCode = newcode
            faculty.Status = 'A'

    firstname = request.POST.get("firstname")
    middlename = request.POST.get("middlename")
    lastname = request.POST.get("lastname")

    if firstname is None:
        firstname = ""
    if middlename is None:
        middlename = ""
    if lastname is None:
        lastname = ""

    faculty.FirstName = firstname
    faculty.MiddleName = middlename
    faculty.LastName = lastname
    faculty.FullName = firstname + " " + middlename + " "+lastname
    faculty.AadharNo = request.POST.get("aadhar")
    faculty.JoiningDate = request.POST.get("joiningdate")
    faculty.Active = True

    if "photo" in request.POST:
        faculty.Photo = request.POST.get("photo")
    faculty.BankAccountNo = request.POST.get("bankaccount")
    faculty.BankCode = request.POST.get("bankcode")
    faculty.CreatedBy = "admin"
    faculty.DateOfBirth = request.POST.get("dob")
    faculty.Education = request.POST.get("education")
    faculty.EmailId = request.POST.get("email")
    faculty.Gender =request.POST.get("gender")
    faculty.IFSC = request.POST.get("ifsc")
    faculty.Mobile = request.POST.get("mobile")
    faculty.Whatsapp = request.POST.get("whatsapp")
    faculty.Salary = request.POST.get("salary")
    faculty.Experience = request.POST.get("experience")
    faculty.NameInBank = request.POST.get("nameinbank")
    faculty.save()

    FacultySubject.objects.filter(FacultyCode = newcode).delete()
    subjects = request.POST.get("subjects").split("^")
    for s in subjects:
        if s != "":
            su = FacultySubject()
            su.FacultyCode = newcode
            su.SubjectId = s
            su.save()

    loginpass = request.POST.get("loginpass")
    if loginpass != "":
        usr = MyUser.objects.filter(UserId = newcode)
        if len(usr) > 0:
            usr = usr[0]
            usr.Password = loginpass
            usr.save()
        else:
            usr = MyUser()
            usr.UserId = newcode
            usr.Password = loginpass
            usr.EmployeeCode = newcode
            usr.UserType = "F"
            usr.save()

            defs = RolePrivilege.objects.filter(RoleCode = "F")
            for d in defs:
                u = UserPrivilege()
                u.UserId = newcode
                u.PrivilegeCode = d.PrivilegeCode
                u.save()

    return JsonResponse({"code":newcode})

@csrf_exempt
def save_staff(request):
    staff = Staff()

    if "staffcode" in request.POST:
        staff = Staff.objects.filter(StaffCode = request.POST.get("staffcode"))
        if len(staff) > 0:
            staff = staff[0]
            newcode = staff.StaffCode
        else:
            staff = Staff()

    firstname = request.POST.get("firstname")
    middlename = request.POST.get("middlename")
    lastname = request.POST.get("lastname")

    if firstname is None:
        firstname = ""
    if middlename is None:
        middlename = ""
    if lastname is None:
        lastname = ""

    staff.FirstName = firstname
    staff.MiddleName = middlename
    staff.LastName = lastname
    staff.FullName = firstname + " " + middlename + " "+lastname
    staff.AadharNo = request.POST.get("aadhar")
    staff.JoiningDate = request.POST.get("joiningdate")
    staff.Active = True
    staff.Photo = request.POST.get("photo")
    staff.BankAccountNo = request.POST.get("bankaccount")
    staff.BankCode = request.POST.get("bankcode")
    staff.CreatedBy = "admin"
    staff.DateOfBirth = request.POST.get("dob")
    staff.Education = request.POST.get("education")
    staff.EmailId = request.POST.get("email")
    staff.Gender =request.POST.get("gender")
    staff.IFSC = request.POST.get("ifsc")
    staff.Mobile = request.POST.get("mobile")
    staff.Whatsapp = request.POST.get("whatsapp")
    staff.Salary = request.POST.get("salary")
    staff.Designation = request.POST.get("designation")
    staff.NameInBank = request.POST.get("nameinbank")
    staff.Status = 'A'
    staff.save()
    newid = staff.id
    newcode = "STF-"+'{:05d}'.format(newid)
    staff.StaffCode = newcode
    staff.save()

    return JsonResponse({"code":newcode})

@csrf_exempt
def save_student(request):
    student = Student()

    if "studentcode" in request.POST:
        student = Student.objects.filter(StudentCode = request.POST.get("studentcode"))
        if len(student) > 0:
            student = student[0]
            newcode = student.StudentCode
        else:
            student = Student()

    firstname = request.POST.get("firstname")
    middlename = request.POST.get("middlename")
    lastname = request.POST.get("lastname")

    if firstname is None:
        firstname = ""
    if middlename is None:
        middlename = ""
    if lastname is None:
        lastname = ""

    student.RollNumber = request.POST.get("rollnumber")
    student.FirstName = firstname
    student.MiddleName = middlename
    student.LastName = lastname
    student.FullName = firstname + " " + middlename + " "+lastname
    student.AadharNo = request.POST.get("aadhar")
    student.Active = True

    if "photo" in request.POST:
        student.Photo = request.POST.get("photo")
    student.CreatedBy = "admin"
    student.DateOfBirth = request.POST.get("dob")
    student.Section = request.POST.get("section")
    student.EmailId = request.POST.get("email")
    student.Gender =request.POST.get("gender")
    student.FatherName = request.POST.get("fathername")
    student.MotherName = request.POST.get("mothername")
    student.Mobile = request.POST.get("mobile")
    student.Whatsapp = request.POST.get("whatsapp")
    student.Class = request.POST.get("class")
    
    if request.POST.get("bus_service") == "true":
        student.BusService = True
        student.BusId = request.POST.get("bus")
    else:
        student.BusService = False
        student.BusId = None
    student.Status = 'A'
    student.save()
    newid = student.id
    newcode = "STD-"+'{:05d}'.format(newid)
    student.StudentCode = newcode
    student.save()
    return JsonResponse({"code":newcode})

@csrf_exempt
def save_member(request):
    comp = request.session["COMPANY"]
    member = LibraryMember()
    mycomp = Company.objects.filter(CompanyCode = comp)
    prefix = "MBR"
    if len(mycomp) > 0:
        mycomp = mycomp[0]
        prefix = mycomp.MemberPrefix

    if "membercode" in request.POST:
        member = LibraryMember.objects.filter(MemberCode = request.POST.get("membercode"))
        if len(member) > 0:
            member = member[0]
            newcode = member.MemberCode
        else:
            member = LibraryMember()

    member.FullName = request.POST.get("fullname")
    member.AadharNo = request.POST.get("aadhar")
    member.Active = True
    member.CompanyCode = comp
    if "photo" in request.POST:
        member.Picture = request.POST.get("photo")
    member.BatchId = request.POST.get("batch")
    member.Email = request.POST.get("email")
    member.Gender =request.POST.get("gender")
    member.Mobile = request.POST.get("mobile")
    member.Whatsapp = request.POST.get("whatsapp")
    member.JoiningDate = request.POST.get("joiningdate")
    member.Status = 'A'
    member.MonthlyFee = request.POST.get("monthlyfee")
    member.save()

    allmembers = LibraryMember.objects.filter(MemberCode__contains = prefix)
    newid = len(allmembers)+1
    newcode = prefix+"-"+'{:05d}'.format(newid)
    member.MemberCode = newcode
    member.save()

    return JsonResponse({"code":newcode})


@csrf_exempt
def get_faculty_list(request):
    status = request.POST.get("status")
    search = request.POST.get("search")
    alldata = None
    with connection.cursor() as cursor:
        cursor.execute("""select f.*,
        (select GROUP_CONCAT(s.Title, ', ') from main_FacultySubject fs
        left join main_Subject s on s.id = fs.SubjectId
        where fs.FacultyCode = f.FacultyCode) as SubjectName
        from main_Faculty f
        where f.FullName like %s and f.status = %s and f.Deleted=0""", ['%'+search+'%', status])

        alldata = parse_curser(cursor)
    return JsonResponse(alldata, safe=False)

@csrf_exempt
def get_staff_list(request):
    status = request.POST.get("status")
    search = request.POST.get("search")
    alldata = None
    with connection.cursor() as cursor:
        cursor.execute("""select f.*, d.Title as designationtitle
        from main_Staff f
        left join main_Designation d on d.id = f.Designation
        where f.FullName like %s and f.status = %s and f.Deleted=0""", ['%'+search+'%', status])

        alldata = parse_curser(cursor)
    return JsonResponse(alldata, safe=False)

@csrf_exempt
def get_student_list(request):
    status = request.POST.get("status")
    search = request.POST.get("search")
    section = request.POST.get("section")
    
    cls = request.POST.get("class")
    alldata = []
    with connection.cursor() as cursor:
        cursor.execute("""select s.*, c.Title as classtitle, ' - ' ||t.Title as SectionTitle from main_Student s
        left join main_Class c on c.id = s.Class
        left join main_Section t on t.id = s.Section
        where s.Deleted = 0 and s.Status = %s and s.FullName like %s and (%s = '' or %s = s.Class) and (%s = '' or %s = s.Section)
        order by s.Class desc, s.RollNumber""",
        [status,"%"+search+"%" ,cls, cls, section, section])

        alldata = parse_curser(cursor)
    return JsonResponse(alldata, safe=False)

@csrf_exempt
def get_member_list(request):
    status = request.POST.get("status")
    search = request.POST.get("search")
    batch = request.POST.get("batch")
    comp = request.session["COMPANY"]

    alldata = []
    with connection.cursor() as cursor:
        cursor.execute("""select s.*, b.Title as batchtitle from main_LibraryMember s
        left join main_Batch b on b.id = s.BatchId
        where s.Deleted = 0 and s.Status = %s and (s.FullName like %s or s.MemberCode like %s) and (%s = '' or %s = s.BatchId)
        and (%s='' or %s=s.CompanyCode)
        order by s.FullName""",
        [status,"%"+search+"%", "%"+search+"%" ,batch, batch, comp, comp])

        alldata = parse_curser(cursor)
    return JsonResponse(alldata, safe=False)

@csrf_exempt
def save_report_card(request):
    classid = request.POST.get("classid")
    sectionid = request.POST.get("sectionid")
    examid = request.POST.get("examid")
    datastring = request.POST.get("datastring")

    datasplit = datastring.split("#")
    for data in datasplit[:-1]:
        sp1 = data.split("*")
        scode = sp1[0]
        markstring = sp1[1]
        marksplit = sp1[1].split(">")

        old = ReportCard.objects.filter(ExamId = examid, StudentCode = scode)
        old.delete()

        for m in marksplit:
            if m != "":
                report = ReportCard()
                report.StudentCode = scode
                report.ExamId = examid
                report.ClassId = classid
                report.SectionId = sectionid
                report.PaperId = m.split("/")[0]
                report.Marks = m.split("/")[1]
                report.save()

    return JsonResponse({"status":"success"})

@csrf_exempt
def get_fee_list(request):
    search = request.POST.get("search")
    sessionid = request.POST.get("sessionid")
    alldata = []
    with connection.cursor() as cursor:
        cursor.execute("""select f.*, s.FullName as StudentName,
          m.Title as ModeTitle, t.Title as TypeTitle, c.Title as classTitle, f.DocumentCode as docoumentcode
        from main_Fee f
        left join main_Student s on s.StudentCode = f.StudentCode
        left join main_Class c on c.id = s.Class
        left join main_PaymentMode m on m.id = f.Mode
        left join main_FeeType t on t.id = f.Type
        where f.Deleted = 0 and s.FullName like %s and f.SessionId=%s""", ['%'+search+'%', sessionid])

        alldata = parse_curser(cursor)
    return JsonResponse(alldata, safe=False)

@csrf_exempt
def member_payment_record(request):
    code = request.POST.get("membercode")
    comp = request.session["COMPANY"]
    record = []
    with connection.cursor() as cursor:
        cursor.execute("""select fi.*,ft.Title as FeeTypeTitle, f.DocumentDate,
        f.TotalAmount as NetAmount, m.title as modetitle, f.Referenceno, f.remarks
        from main_Fee f
        left join main_FeeItem fi on fi.FeeId = f.id
        left join main_FeeType ft on ft.id = fi.Type
        left join main_PaymentMode m on m.id = f.mode
        where f.StudentCode = %s and f.CompanyCode = %s and f.Deleted = 0
        order by cast(replace(fi.monthyear, '-','') as int) desc""", [code, comp])

        record = parse_curser(cursor)

    return JsonResponse(record, safe=False)

@csrf_exempt
def student_payment_record(request):
    code = request.POST.get("studentcode")
    record = []
    with connection.cursor() as cursor:
        cursor.execute("""select fi.*,ft.Title as FeeTypeTitle, f.DocumentDate,
        f.TotalAmount as NetAmount, m.title as modetitle, f.Referenceno, f.remarks
        from main_Fee f
        left join main_FeeItem fi on fi.FeeId = f.id
        left join main_FeeType ft on ft.id = fi.Type
        left join main_PaymentMode m on m.id = f.mode
        where f.StudentCode = %s and f.Deleted = 0
        order by cast(replace(fi.monthyear, '-','') as int) desc""", [code])

        record = parse_curser(cursor)

    return JsonResponse(record, safe=False)

@csrf_exempt
def get_fee_item_detail(request):
    itemid = request.POST.get("itemid")
    detail = []
    with connection.cursor() as cursor:
        cursor.execute("""select fi.*, f.Remarks, f.ReferenceNo, 
        m.Title as ModeTitle, ft.Title as TypeTitle, 
        f.DocumentDate, f.TotalAmount as NetAmount,
        f.Mode
        from main_FeeItem fi
        inner join main_Fee f on f.id = fi.Feeid
        left join main_PaymentMode m on m.id = f.Mode
        left join main_FeeType ft on ft.id = fi.Type
        where fi.id=%s""", [itemid])

        detail = parse_curser(cursor)

    return JsonResponse(detail,safe=False)

@csrf_exempt
def get_member_fee_list(request):
    comp = request.session["COMPANY"]
    batch = request.POST.get("batch")
    #month = request.POST.get("filtermonth")

    search = request.POST.get("search")
    alldata = []
    with connection.cursor() as cursor:
        cursor.execute("""select s.*, c.Title as batchTitle,
        (select fi.monthyear from main_Fee f 
        left join main_FeeItem fi on fi.FeeId = f.id 
        where f.StudentCode = s.MemberCode and f.Deleted = 0
        order by CAST(REPLACE(fi.monthyear, '-','') as int) DESC limit 1) as lastfee
        from main_LibraryMember s
        left join main_Batch c on c.id = s.BatchId
        where s.FullName like %s and
        (%s='' or %s=s.CompanyCode) and (%s='' or %s=s.BatchId) 
        order by FullName""", ['%'+search+'%', comp, comp, batch, batch])

        alldata = parse_curser(cursor)
    return JsonResponse(alldata, safe=False)


@csrf_exempt
def get_faculty_detail(request):
    fcode = request.POST.get("facultycode")
    faculty = []

    with connection.cursor() as cursor:
        cursor.execute("""select f.*, u.password as loginpass from main_Faculty f
            left join main_MyUser u on u.userid = f.FacultyCode 
            where f.facultycode = %s""", [fcode])
        
        faculty = parse_curser(cursor)

    subjects = []
    with connection.cursor() as cursor:
            cursor.execute('''SELECT s.* FROM main_Faculty f
            left join main_FacultySubject fs on f.FacultyCode=fs.FacultyCode
            left join main_Subject s on s.id = fs.SubjectId
            WHERE f.FacultyCode=%s''', [request.POST.get('facultycode')])

            subjects = parse_curser(cursor)
    return JsonResponse({"data":faculty, "subjects":subjects})

@csrf_exempt
def get_staff_detail(request):
    staff = Staff.objects.filter(StaffCode = request.POST.get("staffcode"))
    dt = parse_data(staff)[0]
    return JsonResponse({"data":dt})

@csrf_exempt
def get_expense_detail(request):
    expense = []
    with connection.cursor() as cursor:
        cursor.execute("""select e.*, m.title as modetitle, t.title as typetitle,
                       s.title as sessiontitle 
                       from main_Expense e
                       left join main_PaymentMode m on e.mode = m.id 
                       left join main_ExpenseType t on t.id = e.type
                       left join main_SchoolSession s on s.id = e.sessionid
                       where e.ExpenseCode = %s""", [request.POST.get("expensecode")])

        expense = parse_curser(cursor)
    dt = expense[0]
    return JsonResponse({"data":dt})

@csrf_exempt
def get_paper_detail(request):
    paper = ExamPaper.objects.filter(id = request.POST.get("paperid"))
    alldata = parse_data(paper)[0]

    return JsonResponse(alldata, safe=False)

@csrf_exempt
def get_paper_list(request):
    alldata = None
    with connection.cursor() as cursor:
        cursor.execute("""select p.*, s.Title as Subject, c.Title as Class,
        ' - '||t.Title as Section, e.Title as Exam, (case when p.paperdate < CURRENT_DATE then '1' else '0' end) as Finished
        from main_ExamPaper p
        left join main_Subject s on s.id = p.subjectid
        left join main_Class c on c.id = p.ClassId
        left join main_Section t on t.id = p.SectionId
        left join main_Exam e on e.id = p.ExamId
        where p.Deleted = 0 order by paperdate desc""")

        alldata = parse_curser(cursor)

    return JsonResponse(alldata, safe=False)

@csrf_exempt
def get_expense_list(request):
    search = request.POST.get("search")
    sessionid = request.POST.get("sessionid")
    alldata = []
    with connection.cursor() as cursor:
            cursor.execute('''SELECT e.*, t.Title as typetitle FROM main_Expense e
            left join main_ExpenseType t on e.Type = t.id
            WHERE e.Sessionid = %s and e.Deleted=0 and (e.expensecode like %s or t.Title like %s)''',
            [sessionid,"%"+search+"%", "%"+search+"%"])

            alldata = parse_curser(cursor)
    return JsonResponse(alldata, safe=False)


@csrf_exempt
def get_student_detail(request):
    student = []
    StudentCode = request.POST.get("studentcode")
    with connection.cursor() as cursor:
        cursor.execute("""select s.*, b.Name || ' : ' || b.busno as busname from main_Student s
        left join main_Bus b on b.id = s.busid
        where s.studentcode = %s""", [StudentCode])
        student = parse_curser(cursor)[0]

    return JsonResponse({"data":student})

@csrf_exempt
def get_member_detail(request):
    member = []
    membercode = request.POST.get("membercode")
    with connection.cursor() as cursor:
        cursor.execute("""select s.*, b.title as batchtitle,
        (select mu.Password from main_MyUser mu where mu.UserId = s.MemberCode LIMIT 1)  as loginpassword,
        (select mu.Blocked from main_MyUser mu where mu.UserId = s.MemberCode LIMIT 1)  as blocked
        from main_LibraryMember s
        left join main_Batch b on b.id = s.BatchId
        where s.membercode = %s""", [membercode])
        member = parse_curser(cursor)[0]

    return JsonResponse({"data":member})

@csrf_exempt
def get_student_profile(request):
    student = []
    with connection.cursor() as cursor:
        cursor.execute("""select s.*, c.Title as classtitle, t.Title as sectiontitle
        from main_Student s
        left join main_Class c on c.id = s.Class
        left join main_Section t on t.id = s.Section
        where s.StudentCode = %s""", [request.POST.get("studentcode")])

        student = parse_curser(cursor)
    return JsonResponse({"data":student})

@csrf_exempt
def get_member_profile(request):
    student = []
    with connection.cursor() as cursor:
        cursor.execute("""select s.*, c.Title as batchtitle
        from main_LibraryMember s
        left join main_Batch c on c.id = s.BatchId
        where s.MemberCode = %s""", [request.POST.get("membercode")])

        student = parse_curser(cursor)
    return JsonResponse({"data":student})

@csrf_exempt
def get_fee_detail(request):
    fee = Fee.objects.filter(id = request.POST.get("feeid"))
    feeitem = FeeItem.objects.filter(FeeId = request.POST.get("feeid"))
    dt = parse_data(fee)[0]
    feeitem = parse_data(feeitem)
    return JsonResponse({"data":dt, "items":feeitem})

@csrf_exempt
def get_fee_detail_show(request):
    feeid = request.POST.get("feeid")
    fee = None
    items = None
    with connection.cursor() as cursor:
        cursor.execute('''select f.*, s.FullName, ss.Title as SessionTitle, m.Title as ModeTitle from main_Fee f
        left join main_Student s on s.StudentCode = f.StudentCode
        left join main_SchoolSession ss on ss.id = f.SessionId
        left join main_PaymentMode m on m.id = f.Mode
        where f.id=%s''', [feeid])
        fee = parse_curser(cursor)

    with connection.cursor() as cursor:
        cursor.execute('''select i.*, p.Title as typetitle from main_FeeItem i
        left join main_FeeType p on p.id = i.type
        where i.FeeId=%s''', [feeid])

        items = parse_curser(cursor)

    return JsonResponse({"detail":fee, "items": items})

@csrf_exempt
def get_member_fee_detail_show(request):
    feeid = request.POST.get("feeid")
    fee = None
    items = None
    with connection.cursor() as cursor:
        cursor.execute('''select f.*, s.FullName, m.Title as ModeTitle from main_Fee f
        left join main_LibraryMember s on s.MemberCode = f.StudentCode
        left join main_PaymentMode m on m.id = f.Mode
        where f.id=%s''', [feeid])
        fee = parse_curser(cursor)

    with connection.cursor() as cursor:
        cursor.execute('''select i.*, p.Title as typetitle from main_FeeItem i
        left join main_FeeType p on p.id = i.type
        where i.FeeId=%s''', [feeid])

        items = parse_curser(cursor)

    return JsonResponse({"detail":fee, "items": items})


@csrf_exempt
def save_file(request):
    if "myfile" in request.FILES:
        myfile = request.FILES["myfile"]
        fs = FileSystemStorage(location=django_settings.STATIC_ROOT + '/uploads')
        newname = generate_id(15)+"."+myfile.name.split(".")[-1]
        filename = fs.save(newname, myfile)
        return JsonResponse({"name":newname})
    else:
        return JsonResponse({"name":"Not Found"})

@csrf_exempt
def change_faculty_status(request):
    code = request.POST.get("facultycode")
    status = request.POST.get("status")

    faculty=Faculty.objects.filter(FacultyCode = code)
    if len(faculty) > 0:
        faculty = faculty[0]
        faculty.Status = status
        faculty.save()

    return JsonResponse({"status":"success"})

@csrf_exempt
def change_staff_status(request):
    code = request.POST.get("staffcode")
    status = request.POST.get("status")

    staff = Staff.objects.filter(StaffCode = code)
    if len(staff) > 0:
        staff = staff[0]
        staff.Status = status
        staff.save()

    return JsonResponse({"status":"success"})

@csrf_exempt
def change_student_status(request):
    code = request.POST.get("studentcode")
    status = request.POST.get("status")

    faculty=Student.objects.filter(StudentCode = code)
    if len(faculty) > 0:
        faculty = faculty[0]
        faculty.Status = status
        faculty.save()

    return JsonResponse({"status":"success"})

@csrf_exempt
def change_member_status(request):
    code = request.POST.get("membercode")
    status = request.POST.get("status")
    leavedate = request.POST.get("leavingdate")

    faculty=LibraryMember.objects.filter(MemberCode = code)
    if len(faculty) > 0:
        faculty = faculty[0]
        faculty.Status = status
        if status == 'C':
            faculty.LeavingDate = leavedate
        if status == 'A':
            faculty.LeavingDate = None
        faculty.save()

    return JsonResponse({"status":"success"})

@csrf_exempt
def get_status(request):
    type = request.POST.get("type")
    status = StatusMaster.objects.filter(Type = type)
    comp = request.session["COMPANY"]

    table = None
    data = []
    if type == "FACULTY":
        table = Faculty.objects.filter(Deleted = False)
    elif type == "STUDENT":
        table = Student.objects.filter(Deleted = False)
    elif type == "STAFF":
        table = Staff.objects.filter(Deleted = False)
    elif type == "LIBMEMBER":
        if comp == "":
            table = LibraryMember.objects.filter(Deleted = False)
        else:
            table = LibraryMember.objects.filter(Deleted = False, CompanyCode = comp)

    for s in status:
        count = 0
        for f in table:
            if f.Status == s.StatusCode:
                count+=1

        data.append([s.StatusCode, s.Title, count])

    return JsonResponse({"data":data})

@csrf_exempt
def fee_expense_summary(request):
    month = request.POST.get("month")
    session = request.POST.get("session")

    allfee = []
    with connection.cursor() as cursor:
        cursor.execute("""select SUM(f.totalamount) as TotalFee
        from main_Fee f where strftime('%%Y-%%m', f.DocumentDate)=%s and f.sessionid=%s""", [str(timezone.now())[0:7], session])
        allfee = parse_curser(cursor)

    allexpense = []
    with connection.cursor() as cursor:
        cursor.execute("""select SUM(f.Amount) as TotalExpense
        from main_Expense f where strftime('%%Y-%%m', f.DocumentDate)=%s and f.SessionId=%s""", [str(timezone.now())[0:7],session])
        allexpense = parse_curser(cursor)

    overdue = []
    with connection.cursor() as cursor:
        cursor.execute("""select SUM(c.ClassFee) as overdue from main_Student s
        inner join main_Class c on c.id = s.Class
        where not exists(select f.id from main_Fee f 
        inner join main_FeeItem i on i.Feeid = f.id
        where f.StudentCode = s.StudentCode and i.MonthYear=%s and f.Sessionid=%s)""", [str(timezone.now())[0:7], session])
        overdue = parse_curser(cursor)

    return JsonResponse({"fee":allfee, "expense":allexpense, "overdue":overdue})

@csrf_exempt
def get_member_count(request):
    counts = []
    with connection.cursor() as cursor:
        cursor.execute("""select (select COUNT(m.id) from main_LibraryMember m where m.batchid = b.id and m.Active = 1 and m.Status='A' and m.Deleted = 0) as BatchCount,
        (select COUNT(m.id) from main_LibraryMember m 
        inner join main_Attendance a on m.MemberCode = a.StudentCode 
        where m.batchid = b.id and m.Active = 1 and a.Status = 'P' and strftime('%%Y-%%m-%%d', a.AttendanceDate) = %s and m.Status='A' and m.Deleted = 0) as PresentCount,
        b.Title as BatchTitle
        from main_Batch b 
        where b.CompanyCode=%s""", [str(timezone.now())[0:10],request.session["COMPANY"]])

        counts = parse_curser(cursor)

    return JsonResponse(counts, safe=False)

@csrf_exempt
def get_people_count(request):
    dt = request.POST.get("date")

    studentrecord = []
    facultyrecord = []
    staffrecord = []
    with connection.cursor() as cursor:
        cursor.execute("""select COUNT(s.id) as totalstudent,
        (select COUNT(a.id) from main_Attendance a where strftime('%%Y-%%m-%%d', a.AttendanceDate) = %s and a.Status = 'P') as presentstudent
        from main_Student s
        """, [str(timezone.now())[0:10]])

        studentrecord = parse_curser(cursor)[0]

    with connection.cursor() as cursor:
        cursor.execute("""select COUNT(s.id) as totalfaculty,
        (select COUNT(a.id) from main_FacultyAttendance a where strftime('%%Y-%%m-%%d', a.AttendanceDate) = %s 
        and a.Status = 'P' and SUBSTR(a.FacultyCode, 1, 3) = 'FAC') as presentfaculty
        from main_Faculty s
        """, [str(timezone.now())[0:10]])

        facultyrecord = parse_curser(cursor)[0]

    with connection.cursor() as cursor:
        cursor.execute("""select COUNT(s.id) as totalstaff,
        (select COUNT(a.id) from main_FacultyAttendance a where strftime('%%Y-%%m-%%d', a.AttendanceDate) = %s 
        and a.Status = 'P' and SUBSTR(a.FacultyCode,1,3) = 'STF') as presentstaff
        from main_staff s
        """, [str(timezone.now())[0:10]])

        staffrecord = parse_curser(cursor)[0]
        print(staffrecord)

    return JsonResponse({"faculty": facultyrecord ,"student": studentrecord, "staff": staffrecord})

@csrf_exempt
def load_monthwise_expense(request):
    allexpense = []
    
    with connection.cursor() as cursor:
        cursor.execute('''select *, cast(DocumentDate as varchar(30)) as newdate from main_Expense
        where Deleted = 0''')
        
        expense = parse_curser(cursor)

        for i in range(1,13):
            mn = "2023-"+'{:02d}'.format(i)
            monthdata = 0
            for e in expense:
                if e["newdate"][:7] == mn:
                    am = float(e["amount"])
                    monthdata += am

            allexpense.append(monthdata)

    return JsonResponse({"data":allexpense})

@csrf_exempt
def load_monthwise_library_revenue(request):
    allrevenue = []
    
    with connection.cursor() as cursor:
        cursor.execute("""select COALESCE(fi.TotalAmount, 0) as TotalRevenue,
        fi.MonthYear as newmonth
        from main_Fee f
        inner join main_LibraryMember m on f.StudentCode = m.MemberCode
        inner join main_FeeItem fi on fi.feeid = f.id
        where m.CompanyCode = %s""",
        [request.session["COMPANY"]])

        revenue = parse_curser(cursor)

        for i in range(1,13):
            mn = timezone.now().strftime("%Y")+"-"+'{:02d}'.format(i)
            monthdata = 0
            for e in revenue:
                if e["newmonth"] == mn:
                    am = float(e["totalrevenue"])
                    monthdata += am

            allrevenue.append(monthdata)

    return JsonResponse(allrevenue, safe=False)

@csrf_exempt
def load_member_graph(request):
    members = []
    
    mems = LibraryMember.objects.filter(Deleted = False, CompanyCode = request.session["COMPANY"], Active=True, Status='A')
    cm = int(timezone.now().strftime("%m"))
    year = int(timezone.now().strftime("%Y"))
    for i in range(12):
        lastday = newcalendar.monthrange(year, cm)[1]
        newdate = datetime(year, cm, lastday, 0, 0, 0).date()

        monf = int(str(year)+"{:02d}".format(cm))
        
        ncount = 0
        for m in mems:
            joindate = m.JoiningDate
            leavedate = m.LeavingDate

            if leavedate is None:
                current_datetime = timezone.now()
                leavedate = current_datetime + timedelta(days=1)
                leavedate = leavedate.date()
                leavedate = newdate
            if joindate is None:
                joindate = datetime(2000, 1, 1, 0,0,0).date()

            joined = int(joindate.strftime("%Y%m"))

            if newdate <= leavedate and joined <= monf:
                ncount += 1
        members.append(ncount)

        cm = cm - 1
        if cm == 0:
            cm = 12
            year = year - 1

    return JsonResponse(members, safe=False)

@csrf_exempt
def delete_item(request):
    id = request.POST.get("id")
    table = request.POST.get("table")
    with connection.cursor() as cursor:
        sql = "UPDATE main_{} SET Deleted = 1 WHERE id = %s".format(table)
        cursor.execute(sql, [id])

    return JsonResponse({"status": "success"})

@csrf_exempt
def block_user(request):
    userid = request.POST.get("userid")
    status = request.POST.get("status")

    with connection.cursor() as cursor:
        cursor.execute("""update main_MyUser set Blocked = %s where UserId=%s""", [status, userid])

    return JsonResponse({"status":"success"})

@csrf_exempt
def block_unblock_user(request):
    userid = request.POST.get("userid")
    
    usr = MyUser.objects.filter(UserId = userid)
    if len(usr) > 0:
        usr = usr[0]
        usr.Blocked = not usr.Blocked
        usr.save()
    return JsonResponse({"status":"success"})

@csrf_exempt
def permanent_delete(request):
    id = request.POST.get("id")
    table = request.POST.get("table")

    with connection.cursor() as cursor:
        sql = "delete from main_{} WHERE id = %s".format(table)
        cursor.execute(sql, [id])

    return JsonResponse({"status": "success"})